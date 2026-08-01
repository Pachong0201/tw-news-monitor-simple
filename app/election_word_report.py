"""Build the election analysis report as a real Word document."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)


def _set_style(style, *, size, color, before, after, line_spacing=1.1):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def _section_text(report: dict, key: str) -> str:
    value = report.get(key, "")
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        parts = [value.get("situation", ""), value.get("outlook", "")]
        return "\n\n".join(str(part).strip() for part in parts if str(part).strip())
    return str(value).strip() if value else ""


def _add_text(doc: Document, text: str) -> None:
    if not text:
        doc.add_paragraph("本节暂无可核实内容。")
        return
    for block in text.split("\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)


def build_election_word_report(
    report: dict,
    evidence: dict,
    output_path: str | Path,
    *,
    report_date: str,
) -> Path:
    """Create a standard-business-brief DOCX with deterministic sections."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    _set_style(
        styles["Normal"], size=11, color=RGBColor(0, 0, 0),
        before=0, after=6, line_spacing=1.1,
    )
    _set_style(styles["Heading 1"], size=16, color=BLUE, before=16, after=8)
    _set_style(styles["Heading 2"], size=13, color=BLUE, before=12, after=6)
    _set_style(styles["Heading 3"], size=12, color=DARK_BLUE, before=8, after=4)

    header = section.header.paragraphs[0]
    header.text = "台湾选举观察｜综合分析报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "内部研判资料｜请结合原始证据核验"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("ELECTION INTELLIGENCE BRIEF")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run(report.get("title") or "台湾地方选举态势分析")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    meta_run = meta.add_run(f"报告日期：{report_date}　｜　生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = MUTED

    sections = [
        ("一、总体判断", _section_text(report, "overall_judgment")),
        ("二、台南选情", _section_text(report, "tainan")),
        ("三、新北选情", _section_text(report, "new_taipei")),
        ("四、跨区域比较", _section_text(report, "comparison")),
    ]
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        _add_text(doc, content)

    tainan_facts = evidence.get("tainan_facts", [])
    new_taipei_facts = evidence.get("new_taipei_facts", [])
    all_facts = tainan_facts + new_taipei_facts
    sources = sorted({fact.get("source", "") for fact in all_facts if fact.get("source")})
    doc.add_heading("五、证据摘要", level=1)
    doc.add_paragraph(
        f"本报告共参考 {len(all_facts)} 条事实记录，其中台南 {len(tainan_facts)} 条、"
        f"新北 {len(new_taipei_facts)} 条；涉及 {len(sources)} 个信源。"
    )
    if sources:
        doc.add_paragraph("主要信源：" + "、".join(sources[:12]))
    quality = evidence.get("quality_check", [])
    failed_checks = [item for item in quality if item.get("status") == "fail"]
    doc.add_paragraph(
        "质量检查：" + (f"存在 {len(failed_checks)} 项未通过。" if failed_checks else "全部通过。")
    )

    doc.core_properties.title = report.get("title") or "台湾地方选举态势分析"
    doc.core_properties.subject = "Election analysis report"
    doc.core_properties.author = "tw-news-monitor-simple"
    doc.save(output_path)
    return output_path
