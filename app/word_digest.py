from datetime import datetime
import os
from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor
import html


from .models import Article
from .source_registry import is_official_source, get_source_info, get_official_sources
from .international import display_name, is_international_media
from .international_translation import TranslationResult, translate_article
from .time_utils import TAIPEI

CATEGORY_NAMES = {
    "politics": "政治新闻",
    "economy": "经济新闻",
    "military": "军武",
    "international": "国际新闻",
    "religion": "宗教",
}
CATEGORY_ORDER = ["politics", "economy", "military", "international", "religion"]
CATEGORY_NUMBERS = ["（一）", "（二）", "（三）", "（四）", "（五）"]
HYPERLINK_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

def build_word_digest(
    articles: list[Article],
    output_dir: Path,
    generated_at: datetime | None = None,
    catch_up_urls: set[str] | None = None,
    importance_results: list | None = None,
    international_config: dict | None = None,
    international_coverage: dict[str, list[Article]] | None = None,
    international_translations: dict[str, TranslationResult] | None = None,
) -> Path:
    if not articles:
        raise ValueError("No articles to generate Word digest")
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_at = generated_at or datetime.now()
    
    if catch_up_urls is None:
        catch_up_urls = set()
    fresh_count = len([a for a in articles if a.url not in catch_up_urls])
    catch_up_count = len([a for a in articles if a.url in catch_up_urls])
    official_articles = [a for a in articles if is_official_source(a.source_id)]
    media_articles = [a for a in articles if not is_official_source(a.source_id)]
    total = len(articles)
    official_count = len(official_articles)
    media_count = len(media_articles)

    # 国际媒体层 Phase I：启用配置时，国际媒体文章只进“国际媒体”二级栏目，
    # 不再进入（一）~（五）分类小节；国内“国际新闻”类文章仍进（四）国际新闻。
    intl_enabled = bool(international_config and international_config.get("enabled", False))
    if intl_enabled:
        intl_media_articles = [
            a for a in media_articles
            if is_international_media(a.source_name, international_config)
        ]
        domestic_media_articles = [
            a for a in media_articles
            if not is_international_media(a.source_name, international_config)
        ]
    else:
        intl_media_articles = []
        domestic_media_articles = media_articles
    
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Keep package metadata stable for deterministic isolated reruns.  The
    # report's visible generated time remains the caller-supplied value.
    stable_core_time = generated_at
    if stable_core_time.tzinfo is not None:
        stable_core_time = stable_core_time.astimezone(TAIPEI).replace(tzinfo=None)
    doc.core_properties.author = "python-docx"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.revision = 1
    doc.core_properties.created = stable_core_time
    doc.core_properties.modified = stable_core_time
    
    title = doc.add_heading("台湾新闻监测", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{generated_at.strftime('%Y年%m月%d日 %H:%M')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = meta2.add_run(f"新闻总数：{total}条")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if catch_up_count > 0:
        run3a = meta2.add_run(f"\n正常新闻：{fresh_count}条")
        run3a.font.size = Pt(11)
        run3a.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run3b = meta2.add_run(f"\n补发新闻：{catch_up_count}条")
        run3b.font.size = Pt(11)
        run3b.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run3 = meta2.add_run(f"\n官方信源：{official_count}条")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run4 = meta2.add_run(f"\n新闻媒体：{media_count}条")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    heading_num = 0
    
    if official_articles:
        heading_num += 1
        doc.add_heading(f"{'一二三四五六七八九十'[heading_num-1]}、官方信源", level=1)
        official_by_source: dict[str, list[Article]] = {}
        for a in official_articles:
            official_by_source.setdefault(a.source_id, []).append(a)
        official_sources = get_official_sources()
        sub_idx = 0
        for order, sid, info in official_sources:
            if sid not in official_by_source:
                continue
            sub_idx += 1
            cat_arts = official_by_source[sid]
            cat_arts.sort(key=lambda x: (x.published_at.timestamp() if x.published_at else 0, x.position), reverse=True)
            doc.add_heading(f"{'一二三四五六七八九十'[sub_idx-1]}）{info['display_name']}", level=2)
            for idx, article in enumerate(cat_arts, 1):
                p = doc.add_paragraph()
                if importance_results:
                    _lev = next((r.level for a, r in importance_results if a is article), "")
                else:
                    _lev = ""
                _pfx = "【重大】" if _lev == "critical" else "【重点】" if _lev == "important" else ""
                if article.url in catch_up_urls:
                    display_title = f"{idx}. 【补发】{_pfx}{article.title}"
                else:
                    display_title = f"{idx}. {_pfx}{article.title}"
                run = p.add_run(display_title)
                run.bold = True
                run.font.size = Pt(12)

                if article.summary:
                    p = doc.add_paragraph()
                    run = p.add_run(f"梗概：{article.summary}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    run.italic = True
                
                p = doc.add_paragraph()
                run = p.add_run(f"来源：{article.source_name}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                p = doc.add_paragraph()
                run = p.add_run(f"类型：{info['document_type']}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                if article.published_at:
                    p = doc.add_paragraph()
                    run = p.add_run(f"发布时间：{article.published_at.strftime('%Y-%m-%d %H:%M')}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                p = doc.add_paragraph()
                if article.url in catch_up_urls:
                    p = doc.add_paragraph()
                    run = p.add_run("状态：补发")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                p = doc.add_paragraph()
                _add_hyperlink(p, article.url, article.url)
                doc.add_paragraph()
    
    if media_articles:
        heading_num += 1
        doc.add_heading(f"{'一二三四五六七八九十'[heading_num-1]}、新闻媒体", level=1)
        grouped: dict[str, list[Article]] = {}
        for cat in CATEGORY_ORDER:
            grouped[cat] = []
        for article in domestic_media_articles:
            if article.category in grouped:
                grouped[article.category].append(article)
        for cat_arts in grouped.values():
            cat_arts.sort(key=lambda x: (x.published_at.timestamp() if x.published_at else 0, x.position), reverse=True)
        
        cat_index = 0
        for cat in CATEGORY_ORDER:
            cat_articles = grouped.get(cat, [])
            if not cat_articles:
                continue
            cat_name = CATEGORY_NAMES.get(cat, cat)
            num = CATEGORY_NUMBERS[cat_index]
            cat_index += 1
            doc.add_heading(f"{num}{cat_name}", level=2)
            for idx, article in enumerate(cat_articles, 1):
                p = doc.add_paragraph()
                if importance_results:
                    _lev = next((r.level for a, r in importance_results if a is article), "")
                else:
                    _lev = ""
                _pfx = "【重大】" if _lev == "critical" else "【重点】" if _lev == "important" else ""
                if article.url in catch_up_urls:
                    display_title = f"{idx}. 【补发】{_pfx}{article.title}"
                else:
                    display_title = f"{idx}. {_pfx}{article.title}"
                run = p.add_run(display_title)
                run.bold = True
                run.font.size = Pt(12)

                if article.summary:
                    p = doc.add_paragraph()
                    run = p.add_run(f"梗概：{article.summary}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    run.italic = True
                
                p = doc.add_paragraph()
                run = p.add_run(f"来源：{article.source_name}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                if article.published_at:
                    p = doc.add_paragraph()
                    run = p.add_run(f"发布时间：{article.published_at.strftime('%Y-%m-%d %H:%M')}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                
                p = doc.add_paragraph()
                if article.url in catch_up_urls:
                    p = doc.add_paragraph()
                    run = p.add_run("状态：补发")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                p = doc.add_paragraph()
                _add_hyperlink(p, article.url, article.url)
                doc.add_paragraph()

        # 国际媒体二级栏目（Phase I）：位于分类小节之后，编号接续现有分类小节。
            # 仅当 international_config 启用时渲染；栏目内按重要性、发布时间排序。
        if intl_media_articles:
            if cat_index < len(CATEGORY_NUMBERS):
                intl_num = CATEGORY_NUMBERS[cat_index]
            else:
                intl_num = f"（{'一二三四五六七八九十'[cat_index]}）"
            doc.add_heading(f"{intl_num}国际媒体", level=2)
            def _importance_order(article: Article) -> int:
                if importance_results:
                    result = next((r for a, r in importance_results if a is article), None)
                    if result is not None:
                        return {"critical": 0, "important": 1, "normal": 2}.get(result.level, 2)
                return 2

            intl_sorted = sorted(
                intl_media_articles,
                key=lambda x: (
                    _importance_order(x),
                    -(x.published_at.timestamp() if x.published_at else 0),
                    -x.position,
                ),
            )
            for idx, article in enumerate(intl_sorted, 1):
                p = doc.add_paragraph()
                if importance_results:
                    _lev = next((r.level for a, r in importance_results if a is article), "")
                else:
                    _lev = ""
                _pfx = "【重大】" if _lev == "critical" else "【重点】" if _lev == "important" else ""
                translation = (international_translations or {}).get(article.url)
                if translation is None:
                    # Missing mapping is a legal, metadata-only fallback; it
                    # must never look like a real Chinese translation.
                    translation = translate_article(article, translator=None)
                # A fallback must not relabel an English title as Chinese.
                # The canonical English title is rendered in the dedicated
                # field below; only a real translator gets a title line.
                if translation.status == "translated" and translation.cn_title:
                    chinese_title = translation.cn_title
                    if article.url in catch_up_urls:
                        display_title = f"{idx}. 【补发】{_pfx}{chinese_title}"
                    else:
                        display_title = f"{idx}. {_pfx}{chinese_title}"
                    run = p.add_run(display_title)
                    run.bold = True
                    run.font.size = Pt(12)

                # 来源：中文展示名（英文 canonical 名），英文标题原样输出不做翻译
                p = doc.add_paragraph()
                run = p.add_run(
                    f"来源：{display_name(article.source_name, international_config)}"
                    f"（{article.source_name}）"
                )
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                p = doc.add_paragraph()
                run = p.add_run(f"英文原题：{article.title}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                # LLM summaries in this pipeline are Chinese.  Only an RSS/
                # metadata teaser (or an unspecified legacy value) may be
                # presented as the publisher's English summary.
                english_summary = (
                    article.summary
                    if article.summary_source in {None, "rss", "meta"}
                    else None
                )
                p = doc.add_paragraph()
                run = p.add_run(
                    f"英文摘要：{english_summary or '未提供合法英文摘要。'}"
                )
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                run.italic = True

                if article.published_at:
                    published_at = article.published_at
                    if published_at.tzinfo is not None:
                        published_at = published_at.astimezone(TAIPEI)
                    p = doc.add_paragraph()
                    run = p.add_run(f"发布时间（Asia/Taipei）：{published_at.strftime('%Y-%m-%d %H:%M')}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                if translation.status == "translated":
                    translated_summary = translation.cn_summary or "未提供中文摘要。"
                else:
                    translated_summary = ""
                if "系统判断：" in translated_summary:
                    facts, judgment = translated_summary.split("系统判断：", 1)
                    facts = facts.replace("媒体报道事实：", "").strip(" ：")
                    judgment = judgment.strip()
                    p = doc.add_paragraph()
                    run = p.add_run(f"中文摘要：{facts}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    run.italic = True
                    p = doc.add_paragraph()
                    run = p.add_run(f"系统判断：{judgment}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                elif translated_summary:
                    p = doc.add_paragraph()
                    run = p.add_run(f"中文摘要：{translated_summary}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    run.italic = True
                if translation.status == "fallback" and translation.limitation:
                    p = doc.add_paragraph()
                    run = p.add_run(f"限制：{translation.limitation}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

                p = doc.add_paragraph()
                if article.url in catch_up_urls:
                    p = doc.add_paragraph()
                    run = p.add_run("状态：补发")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                p = doc.add_paragraph()
                _add_hyperlink(p, article.url, article.url)

                # 同一事件的其他国际媒体报道（canonical 且有 coverage 时追加一行）
                members = (international_coverage or {}).get(article.url) or []
                others = [m for m in members if m.url != article.url]
                if others:
                    names = [
                        display_name(m.source_name, international_config) for m in others
                    ]
                    tail = "等" if len(names) > 1 else ""
                    p = doc.add_paragraph()
                    run = p.add_run(f"另据{'、'.join(names)}{tail}报道同一事件")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    p = doc.add_paragraph()
                    run = p.add_run("Coverage 原文链接：")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    for link_index, member in enumerate(others):
                        if link_index:
                            p.add_run("；")
                        _add_hyperlink(p, member.url, member.url)
                doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本简报由自动新闻监测程序生成，具体内容以原媒体报道为准。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    
    filename = f"台湾新闻监测_{generated_at.strftime('%Y-%m-%d_%H%M')}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    _normalise_docx_package(output_path)
    return output_path


def _normalise_docx_package(path: Path) -> None:
    """Rewrite ZIP metadata so identical input produces identical bytes."""

    temp_path: Path | None = None
    try:
        fd, raw_temp = tempfile.mkstemp(
            prefix=f".{path.name}.", suffix=".tmp", dir=str(path.parent)
        )
        os.close(fd)
        temp_path = Path(raw_temp)
        with ZipFile(path, "r") as source, ZipFile(temp_path, "w") as target:
            for item in sorted(source.infolist(), key=lambda entry: entry.filename):
                info = ZipInfo(item.filename, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = item.compress_type
                info.comment = item.comment
                info.create_system = item.create_system
                info.external_attr = item.external_attr
                info.extra = item.extra
                info.flag_bits = item.flag_bits
                target.writestr(info, source.read(item.filename))
        os.replace(temp_path, path)
        temp_path = None
    finally:
        if temp_path is not None:
            try:
                temp_path.unlink()
            except OSError:
                pass

def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_REL_TYPE, is_external=True)
    safe_text = html.escape(text)
    hyperlink_xml = (
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<w:r><w:rPr><w:color w:val=\"0563C1\"/><w:u w:val=\"single\"/><w:sz w:val=\"20\"/></w:rPr>"
        f"<w:t xml:space=\"preserve\">{safe_text}</w:t></w:r></w:hyperlink>"
    )
    hyperlink = parse_xml(hyperlink_xml)
    paragraph._p.append(hyperlink)
