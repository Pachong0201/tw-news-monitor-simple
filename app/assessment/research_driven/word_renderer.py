"""最终研判文章 Word 渲染器（research-driven 路径）。

Word 正文就是最终文章本身：标题 + 报告周期信息栏 + 文章正文。
不含 event_id / source_id / claim_id / validator JSON。
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt


def _set_run_font(run, *, name: str = "宋体", size: float = 12, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), name)


def _add_page_number_footer(doc: Document, report_id: str) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("台南选情智能研判系统　报告ID：")
    _set_run_font(run, name="宋体", size=9)
    run2 = para.add_run(report_id)
    _set_run_font(run2, name="宋体", size=9)
    run3 = para.add_run("　第 ")
    _set_run_font(run3, name="宋体", size=9)
    fld_begin = parse_xml('<w:fldChar {ns} w:fldCharType="begin"/>'.format(ns=nsdecls("w")))
    instr = parse_xml('<w:instrText {ns} xml:space="preserve"> PAGE </w:instrText>'.format(ns=nsdecls("w")))
    fld_end = parse_xml('<w:fldChar {ns} w:fldCharType="end"/>'.format(ns=nsdecls("w")))
    run_page = para.add_run()
    run_page._element.append(fld_begin)
    run_page._element.append(instr)
    run_page._element.append(fld_end)
    _set_run_font(run_page, name="宋体", size=9)
    run4 = para.add_run(" 页")
    _set_run_font(run4, name="宋体", size=9)


def word_filename(period_start: str, period_end: str) -> str:
    return f"台南选情研判_{period_start}至{period_end}.docx"


def render_article_word(
    *,
    title: str,
    body_markdown: str,
    output_dir: Path,
    period_start: str,
    period_end: str,
    facts_cutoff: str,
    poll_cutoff: str,
    report_id: str,
    model: str = "",
    generated_at: datetime | None = None,
) -> dict[str, Any]:
    generated_at = generated_at or datetime.now()
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = word_filename(period_start, period_end)
    docx_path = output_dir / filename

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    _set_run_font(run, name="黑体", size=16, bold=True)
    title_para.paragraph_format.space_after = Pt(6)

    # 报告信息栏
    info = {
        "报告周期": f"{period_start} 至 {period_end}",
        "生成时间": generated_at.strftime("%Y-%m-%d %H:%M:%S"),
        "事实截止日": facts_cutoff or "未披露",
        "民调截止日": poll_cutoff or "未披露",
        "生成模型": model or "未披露",
    }
    info_table = doc.add_table(rows=0, cols=1)
    info_table.style = "Table Grid"
    for key, value in info.items():
        row = info_table.add_row()
        cell = row.cells[0]
        cell.text = ""
        r = cell.paragraphs[0].add_run(f"{key}：{value}")
        _set_run_font(r, name="黑体", size=10.5, bold=True)
    doc.add_paragraph()

    # 正文（简单 Markdown 解析：## 标题、- 列表、普通段落）
    for line in body_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            heading = doc.add_paragraph()
            hr = heading.add_run(stripped[2:].strip())
            _set_run_font(hr, name="黑体", size=15, bold=True)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            heading = doc.add_paragraph()
            hr = heading.add_run(stripped[3:].strip())
            _set_run_font(hr, name="黑体", size=13, bold=True)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("### "):
            heading = doc.add_paragraph()
            hr = heading.add_run(stripped[4:].strip())
            _set_run_font(hr, name="黑体", size=12, bold=True)
            heading.paragraph_format.space_before = Pt(6)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.line_spacing = 1.5
            pr = para.add_run(stripped[2:].strip())
            _set_run_font(pr, name="宋体", size=12)
        else:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.line_spacing = 1.5
            pr = para.add_run(stripped)
            _set_run_font(pr, name="宋体", size=12)

    _add_page_number_footer(doc, report_id)
    doc.save(docx_path)

    return {
        "docx_path": str(docx_path),
        "filename": filename,
        "docx_size_bytes": docx_path.stat().st_size,
        "generated_at": generated_at.isoformat(),
    }
