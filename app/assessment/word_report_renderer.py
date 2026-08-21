"""Word 报告渲染器：只从已验证 structured_report_final.json 生成正文。"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


SECTION_ORDER = [
    "一、总体判断",
    "二、本期关键变化",
    "三、陈亭妃整合进展",
    "四、谢龙介组织及竞选动作",
    "五、蓝白合作变化",
    "六、民调与治理议题",
    "七、未来半月走势",
    "八、证据限制",
]

V2_SECTION_ORDER = [
    "一、结论摘要",
    "二、核心研判",
    "三、数据限制与事实附录",
]

DRAFT_LABEL = "【数据不完整草稿】"
DATE_RE = __import__("re").compile(r"\d{4}-\d{2}-\d{2}")


def _disclosure_texts(structured_report: dict) -> list[str]:
    claims = {c.get("claim_id"): c for c in (structured_report.get("claims") or [])}
    texts: list[str] = []
    seen: set[str] = set()
    for cid in structured_report.get("required_disclosures") or []:
        claim = claims.get(cid)
        text = (claim or {}).get("claim_text")
        if text and text not in seen:
            texts.append(str(text))
            seen.add(text)
    for claim in structured_report.get("claims") or []:
        if claim.get("claim_type") == "data_disclosure" and claim.get("claim_text") not in seen:
            texts.append(str(claim["claim_text"]))
            seen.add(str(claim["claim_text"]))
    return texts


def _cutoff_from_disclosures(texts: list[str], keyword: str) -> str:
    for text in texts:
        if keyword in text:
            match = DATE_RE.search(text)
            if match:
                return match.group(0)
    return ""


def _set_run_font(run, *, name: str = "宋体", size: float = 12, bold: bool = False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), name)


def _add_page_number_footer(doc: Document, report_id: str) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("台南市长选情辅助研判系统　报告ID：")
    _set_run_font(run, name="宋体", size=9)
    run2 = para.add_run(report_id)
    _set_run_font(run2, name="宋体", size=9)
    run3 = para.add_run("　第 ")
    _set_run_font(run3, name="宋体", size=9)
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_page = para.add_run()
    run_page._element.append(fld_begin)
    run_page._element.append(instr)
    run_page._element.append(fld_end)
    _set_run_font(run_page, name="宋体", size=9)
    run4 = para.add_run(" 页")
    _set_run_font(run4, name="宋体", size=9)


def render_word_report(
    structured_report: dict,
    *,
    output_dir: Path,
    mode: str = "development",
    generation_validation: dict | None = None,
    manifest: dict | None = None,
    generated_at: datetime | None = None,
) -> dict:
    """生成 Word 报告并返回产物信息。

    仅接受 report_status in (generated, repaired) 且通过校验的报告。
    v2.0（研判单元契约）与 v1.1（历史契约）均支持。
    """
    status = structured_report.get("report_status")
    if status not in ("generated", "repaired"):
        raise ValueError(f"report_status={status} 不允许生成 Word")
    version = str(structured_report.get("schema_version") or "")
    if version == "2.0":
        return _render_v2_word(
            structured_report,
            output_dir=output_dir,
            mode=mode,
            manifest=manifest,
            generated_at=generated_at,
        )
    return _render_v1_word(
        structured_report,
        output_dir=output_dir,
        mode=mode,
        manifest=manifest,
        generated_at=generated_at,
    )


def _render_v1_word(
    structured_report: dict,
    *,
    output_dir: Path,
    mode: str,
    manifest: dict | None,
    generated_at: datetime | None,
) -> dict:
    status = structured_report.get("report_status")
    generation_mode = structured_report.get("generation_mode")
    report_period = structured_report.get("report_period") or {}
    period_start = str(report_period.get("period_start") or "")
    period_end = str(report_period.get("period_end") or "")
    disclosure_texts = _disclosure_texts(structured_report)
    facts_cutoff = _cutoff_from_disclosures(disclosure_texts, "事实")
    poll_cutoff = _cutoff_from_disclosures(disclosure_texts, "民调")
    dc = structured_report.get("data_context") or {}
    facts_cutoff = dc.get("facts_cutoff") or facts_cutoff
    poll_cutoff = dc.get("poll_cutoff") or poll_cutoff
    active_snapshot_id = dc.get("active_snapshot_id") or ""
    coverage_version = dc.get("coverage_version") or ""
    uncovered_dates = dc.get("uncovered_date_range") or []
    if generation_mode == "draft_with_data_gap":
        filename = f"台南市长选情半月研判_数据不完整草稿_{period_start.replace('-', '')}-{period_end.replace('-', '')}.docx"
    else:
        filename = f"台南市长选情半月研判_{period_start.replace('-', '')}-{period_end.replace('-', '')}.docx"

    generated_at = generated_at or datetime.now()
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / filename

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(structured_report.get("title") or "")
    _set_run_font(run, name="黑体", size=18, bold=True)
    title.paragraph_format.space_after = Pt(6)

    if generation_mode == "draft_with_data_gap":
        draft = doc.add_paragraph()
        draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = draft.add_run(DRAFT_LABEL)
        _set_run_font(dr, name="黑体", size=14, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

    # 报告信息栏
    info = {
        "报告周期": f"{period_start} 至 {period_end}",
        "生成时间": generated_at.strftime("%Y-%m-%d %H:%M:%S"),
        "运行模式": mode,
        "报告状态": status,
        "事实截止日": facts_cutoff or "未披露",
        "民调截止日": poll_cutoff or "未披露",
        "当前快照": active_snapshot_id or "未披露",
        "覆盖版本": coverage_version or "未披露",
        "生成Provider": str((manifest or {}).get("provider") or ""),
        "生成模型": str((manifest or {}).get("model") or ""),
    }
    if generation_mode == "draft_with_data_gap":
        info["未覆盖日期范围"] = "、".join(str(x) for x in uncovered_dates) or "未披露"
        info["本报告不得作为完整周期正式报告"] = "是"
    info_table = doc.add_table(rows=0, cols=1)
    info_table.style = "Table Grid"
    for key, value in info.items():
        row = info_table.add_row()
        cell = row.cells[0]
        cell.text = ""
        r = cell.paragraphs[0].add_run(f"{key}：{value}")
        _set_run_font(r, name="黑体", size=10.5, bold=True)
    doc.add_paragraph()

    # 正文
    claims = {c.get("claim_id"): c for c in (structured_report.get("claims") or [])}
    rendered_claim_ids: list[str] = []
    for section_item in structured_report.get("sections") or []:
        heading = doc.add_paragraph()
        hr = heading.add_run(section_item.get("heading") or "")
        _set_run_font(hr, name="黑体", size=14, bold=True)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        for cid in section_item.get("claim_ids") or []:
            claim = claims.get(cid)
            if not claim:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.line_spacing = 1.5
            cr = para.add_run(claim.get("claim_text") or "")
            _set_run_font(cr, name="宋体", size=12)
            if cid not in rendered_claim_ids:
                rendered_claim_ids.append(cid)

    # 证据附录
    _append_evidence_appendix(doc, structured_report)

    # 数据说明
    data_heading = doc.add_paragraph()
    dr = data_heading.add_run("数据说明")
    _set_run_font(dr, name="黑体", size=14, bold=True)
    for text in disclosure_texts:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        rr = para.add_run(str(text))
        _set_run_font(rr, name="宋体", size=11)
    stats = doc.add_paragraph()
    sr = stats.add_run(
        f"事实截止日：{facts_cutoff or '未披露'}　民调截止日：{poll_cutoff or '未披露'}　"
        f"生成模式：{generation_mode}"
    )
    _set_run_font(sr, name="宋体", size=10.5)

    _add_page_number_footer(doc, str(structured_report.get("report_id") or ""))
    doc.save(docx_path)

    return {
        "docx_path": str(docx_path),
        "filename": filename,
        "report_mode": generation_mode,
        "report_status": status,
        "section_count": len(structured_report.get("sections") or []),
        "claim_count": len(structured_report.get("claims") or []),
        "rendered_claim_count": len(rendered_claim_ids),
        "rendered_claim_ids": rendered_claim_ids,
        "docx_size_bytes": docx_path.stat().st_size,
        "generated_at": generated_at.isoformat(),
    }


def _render_v2_word(
    structured_report: dict,
    *,
    output_dir: Path,
    mode: str,
    manifest: dict | None,
    generated_at: datetime | None,
) -> dict:
    """v2.0 研判单元契约渲染：结论摘要 -> 核心研判 -> 数据限制与事实附录。"""
    status = structured_report.get("report_status")
    generation_mode = structured_report.get("generation_mode")
    report_period = structured_report.get("report_period") or {}
    period_start = str(report_period.get("period_start") or "")
    period_end = str(report_period.get("period_end") or "")
    disclosure_texts = _disclosure_texts_v2(structured_report)
    dc = structured_report.get("data_context") or {}
    facts_cutoff = dc.get("facts_cutoff") or ""
    poll_cutoff = dc.get("poll_cutoff") or ""
    active_snapshot_id = dc.get("active_snapshot_id") or ""
    coverage_version = dc.get("coverage_version") or ""
    uncovered_dates = dc.get("uncovered_date_range") or []
    if generation_mode == "draft_with_data_gap":
        filename = f"台南市长选情半月研判_数据不完整草稿_{period_start.replace('-', '')}-{period_end.replace('-', '')}.docx"
    else:
        filename = f"台南市长选情半月研判_{period_start.replace('-', '')}-{period_end.replace('-', '')}.docx"

    generated_at = generated_at or datetime.now()
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / filename

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(structured_report.get("title") or "")
    _set_run_font(run, name="黑体", size=18, bold=True)
    title.paragraph_format.space_after = Pt(6)

    if generation_mode == "draft_with_data_gap":
        draft = doc.add_paragraph()
        draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = draft.add_run(DRAFT_LABEL)
        _set_run_font(dr, name="黑体", size=14, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

    # 报告信息栏
    info = {
        "报告周期": f"{period_start} 至 {period_end}",
        "生成时间": generated_at.strftime("%Y-%m-%d %H:%M:%S"),
        "运行模式": mode,
        "报告状态": status,
        "事实截止日": facts_cutoff or "未披露",
        "民调截止日": poll_cutoff or "未披露",
        "当前快照": active_snapshot_id or "未披露",
        "覆盖版本": coverage_version or "未披露",
        "生成Provider": str((manifest or {}).get("provider") or ""),
        "生成模型": str((manifest or {}).get("model") or ""),
    }
    if generation_mode == "draft_with_data_gap":
        info["未覆盖日期范围"] = "、".join(str(x) for x in uncovered_dates) or "未披露"
        info["本报告不得作为完整周期正式报告"] = "是"
    info_table = doc.add_table(rows=0, cols=1)
    info_table.style = "Table Grid"
    for key, value in info.items():
        row = info_table.add_row()
        cell = row.cells[0]
        cell.text = ""
        r = cell.paragraphs[0].add_run(f"{key}：{value}")
        _set_run_font(r, name="黑体", size=10.5, bold=True)
    doc.add_paragraph()

    # 一、结论摘要（首屏正文）
    _add_v2_heading(doc, "一、结论摘要")
    for item in structured_report.get("conclusion_summary") or []:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.line_spacing = 1.5
        cr = para.add_run(
            f"{item.get('judgment')}（置信度：{item.get('confidence')}）"
        )
        _set_run_font(cr, name="宋体", size=12)

    # 二、核心研判
    _add_v2_heading(doc, "二、核心研判")
    for index, assessment in enumerate(
        structured_report.get("core_assessments") or [], 1
    ):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.line_spacing = 1.5
        cr = para.add_run(f"研判{index}")
        _set_run_font(cr, name="黑体", size=13, bold=True)

        _add_v2_label_paragraph(doc, "判断", assessment.get("judgment") or "")
        for ev in assessment.get("evidence_items") or []:
            _add_v2_label_paragraph(
                doc,
                "证据",
                f"{ev.get('evidence_date')} {ev.get('evidence_summary')}（{ev.get('evidence_id')}）",
            )
        _add_v2_label_paragraph(doc, "推理", assessment.get("reasoning") or "")
        _add_v2_label_paragraph(
            doc, "反证/限制", assessment.get("falsifiers_or_limits") or ""
        )
        indicators = "；".join(assessment.get("watch_indicators") or [])
        _add_v2_label_paragraph(
            doc,
            "置信度与观察指标",
            f"{assessment.get('confidence')}；{indicators}",
        )

    # 三、数据限制与事实附录
    _add_v2_heading(doc, "三、数据限制与事实附录")
    for item in structured_report.get("appendix") or []:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.line_spacing = 1.5
        cr = para.add_run(f"[{item.get('item_type')}] {item.get('item_text')}")
        _set_run_font(cr, name="宋体", size=11)

    # 证据附录（派生 claim 表）
    _append_evidence_appendix_v2(doc, structured_report)

    # 数据说明
    data_heading = doc.add_paragraph()
    dr = data_heading.add_run("数据说明")
    _set_run_font(dr, name="黑体", size=14, bold=True)
    for text in disclosure_texts:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        rr = para.add_run(str(text))
        _set_run_font(rr, name="宋体", size=11)
    stats = doc.add_paragraph()
    sr = stats.add_run(
        f"事实截止日：{facts_cutoff or '未披露'}　民调截止日：{poll_cutoff or '未披露'}　"
        f"生成模式：{generation_mode}"
    )
    _set_run_font(sr, name="宋体", size=10.5)

    _add_page_number_footer(doc, str(structured_report.get("report_id") or ""))
    doc.save(docx_path)

    return {
        "docx_path": str(docx_path),
        "filename": filename,
        "report_mode": generation_mode,
        "report_status": status,
        "section_count": len(V2_SECTION_ORDER),
        "claim_count": len(structured_report.get("claims") or []),
        "rendered_claim_count": len(structured_report.get("claims") or []),
        "rendered_claim_ids": [
            c.get("claim_id") for c in (structured_report.get("claims") or [])
        ],
        "docx_size_bytes": docx_path.stat().st_size,
        "generated_at": generated_at.isoformat(),
    }


def _add_v2_heading(doc: Document, heading: str) -> None:
    para = doc.add_paragraph()
    hr = para.add_run(heading)
    _set_run_font(hr, name="黑体", size=14, bold=True)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)


def _add_v2_label_paragraph(doc: Document, label: str, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    lr = para.add_run(f"{label}：")
    _set_run_font(lr, name="黑体", size=12, bold=True)
    rr = para.add_run(text)
    _set_run_font(rr, name="宋体", size=12)


def _append_evidence_appendix(doc: Document, structured_report: dict) -> None:
    appendix_heading = doc.add_paragraph()
    ar = appendix_heading.add_run("证据附录")
    _set_run_font(ar, name="黑体", size=14, bold=True)
    appendix_heading.paragraph_format.space_before = Pt(12)
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "Claim ID", "Claim类型", "Claim内容", "event_id", "poll_id",
        "source_id", "snapshot dimension", "gap_id", "confidence", "limitations",
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        hr = cell.paragraphs[0].add_run(header)
        _set_run_font(hr, name="黑体", size=9, bold=True)
    for claim in structured_report.get("claims") or []:
        row = table.add_row()
        values = [
            claim.get("claim_id") or "",
            claim.get("claim_type") or "",
            claim.get("claim_text") or "",
            "、".join(claim.get("supporting_event_ids") or []),
            "、".join(claim.get("supporting_poll_ids") or []),
            "、".join(claim.get("supporting_source_ids") or []),
            "、".join(claim.get("supporting_snapshot_dimensions") or []),
            "、".join(claim.get("supporting_gap_ids") or []),
            claim.get("confidence") or "",
            "；".join(claim.get("limitations") or []),
        ]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(str(value))
            _set_run_font(rr, name="宋体", size=9)


def _append_evidence_appendix_v2(doc: Document, structured_report: dict) -> None:
    claims = structured_report.get("claims") or []
    if not claims:
        # 未富化（原始 v2 模型输出）时按结构派生一份展示用 claim 表。
        from .claim_evidence_validator import build_evidence_context
        from .report_structure_validator import derive_claims_and_sections

        ctx = build_evidence_context(
            structured_report.get("data_context") or {}, evidence_pack=None, config={}
        )
        claims, _ = derive_claims_and_sections(structured_report, ctx)
    appendix_heading = doc.add_paragraph()
    ar = appendix_heading.add_run("证据附录")
    _set_run_font(ar, name="黑体", size=14, bold=True)
    appendix_heading.paragraph_format.space_before = Pt(12)
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "Claim ID", "Claim类型", "Claim内容", "event_id", "poll_id",
        "source_id", "snapshot dimension", "gap_id", "confidence", "limitations",
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        hr = cell.paragraphs[0].add_run(header)
        _set_run_font(hr, name="黑体", size=9, bold=True)
    for claim in claims:
        row = table.add_row()
        values = [
            claim.get("claim_id") or "",
            claim.get("claim_type") or "",
            claim.get("claim_text") or "",
            "、".join(claim.get("supporting_event_ids") or []),
            "、".join(claim.get("supporting_poll_ids") or []),
            "、".join(claim.get("supporting_source_ids") or []),
            "、".join(claim.get("supporting_snapshot_dimensions") or []),
            "、".join(claim.get("supporting_gap_ids") or []),
            claim.get("confidence") or "",
            "；".join(claim.get("limitations") or []),
        ]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(str(value))
            _set_run_font(rr, name="宋体", size=9)


def _disclosure_texts_v2(structured_report: dict) -> list[str]:
    texts: list[str] = []
    seen: set[str] = set()
    for text in structured_report.get("required_disclosures") or []:
        text = str(text).strip()
        if text and text not in seen:
            texts.append(text)
            seen.add(text)
    for item in structured_report.get("appendix") or []:
        if item.get("item_type") == "disclosure":
            text = str(item.get("item_text") or "").strip()
            if text and text not in seen:
                texts.append(text)
                seen.add(text)
    return texts


def extract_word_text(path: Path) -> str:
    """提取 Word 全部可见文本（段落 + 表格），用于业务幂等与校验。"""
    import re

    doc = Document(path)
    parts: list[str] = []
    timestamp_re = re.compile(
        r"^(?:\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}$"
        r"|[^：\n]+：\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}$)"
    )
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not timestamp_re.match(text):
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and not timestamp_re.match(text):
                    parts.append(text)
    return "\n".join(parts)


def extract_word_body(path: Path) -> str:
    """提取正文部分文本（到证据附录之前），用于 claim 顺序校验。"""
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "证据附录":
            break
        if text:
            parts.append(text)
    return "\n".join(parts)
