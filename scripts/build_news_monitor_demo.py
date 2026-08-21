"""Build a non-live Word demonstration of the Taiwan news monitor.

This script reads only the source configuration.  Its articles are explicitly
labelled demonstrations, so it never collects news, writes the news database,
or sends a notification.
"""

from __future__ import annotations

import argparse
from collections import Counter
from datetime import date
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCES_PATH = ROOT / "config" / "sources.yaml"
REPORTS_DIR = ROOT / "data" / "reports"
DEMO_NOTICE = "演示样例，非实时监测结论"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF6DC"
TABLE_WIDTH_DXA = 9360


def enabled_status(source: dict) -> str:
    if source.get("enabled", True) is False:
        return "保留但未启用"
    if source.get("type") == "wsj_newsletter":
        return "已启用（需邮箱授权）"
    if source.get("type") == "newsletter" or "newsletter" in str(source.get("id", "")):
        return "已启用（需邮箱授权）"
    return "已启用"


def source_group(source: dict) -> str:
    source_id = str(source.get("id", ""))
    name = str(source.get("name", ""))
    source_type = str(source.get("type", ""))
    if source_type == "president_json" or name == "行政院":
        return "官方来源"
    if source_id == "ltn_defense":
        return "军武来源"
    if source_type in {"reuters", "ft_alphaville", "wsj_rss"}:
        return "国际媒体"
    if "newsletter" in source_id:
        return "授权通讯源"
    if source_id in {"voice_of_tibet_cn", "zaobao_international"}:
        return "区域中文媒体"
    return "台湾媒体"


def collector_label(source: dict) -> str:
    labels = {
        "rss": "RSS / Atom",
        "udn": "HTML 列表页",
        "ebc": "HTML 列表页",
        "cna_list_html": "HTML 列表页",
        "ltn_rss": "RSS",
        "president_json": "官方 JSON API",
        "newtalk_rss": "RSS",
        "zaobao": "Google News Sitemap",
        "reuters": "官方 Sitemap（仅元数据）",
        "ft_alphaville": "官方 RSS（仅导语）",
        "wsj_rss": "官方 RSS（冻结保留）",
        "wsj_newsletter": "授权邮箱 Newsletter",
        "bloomberg_newsletter": "授权邮箱 Newsletter",
    }
    return labels.get(str(source.get("type", "")), str(source.get("type", "未知")))


def category_label(source: dict) -> str:
    category = str(source.get("category") or source.get("default_category") or "按内容归类")
    return {"politics": "政治", "economy": "财经", "international": "国际", "military": "军武"}.get(category, category)


def load_source_rows(path: Path = SOURCES_PATH) -> list[dict]:
    config = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    rows: list[dict] = []
    for source in config.get("sources", []):
        rows.append(
            {
                "group": source_group(source),
                "name": str(source.get("name", "未命名来源")),
                "id": str(source.get("id", "")),
                "categories": category_label(source),
                "collector": collector_label(source),
                "status": enabled_status(source),
            }
        )
    order = {"官方来源": 0, "台湾媒体": 1, "军武来源": 2, "区域中文媒体": 3, "国际媒体": 4, "授权通讯源": 5}
    return sorted(rows, key=lambda item: (order.get(item["group"], 99), item["name"], item["id"]))


def demo_articles() -> dict[str, list[dict]]:
    return {
        "政治": [
            {"source": "总统府（演示）", "title": "政府公布韧性治理说明会安排", "summary": "演示系统会保留官方来源、发布时间与原始链接，并以完整句呈现要点。"},
            {"source": "中央社（演示）", "title": "立法议题进入跨党派协商程序", "summary": "系统将同类政治新闻统一入库、去重，并依据时效与重要度排序。"},
        ],
        "财经": [
            {"source": "自由時報（演示）", "title": "产业供应链企业发布季度营运展望", "summary": "财经栏目会过滤与监测目标无关的社会资讯，并保留可追溯的新闻来源。"},
            {"source": "联合新闻网（演示）", "title": "出口订单与半导体投资议题受市场关注", "summary": "系统可将 RSS 导语或授权摘要整理为不截断的中文梗概。"},
        ],
        "军武": [
            {"source": "自由时报·军武（演示）", "title": "防务动态与训练整备议题进入监测清单", "summary": "军武频道作为独立来源采集，并按军武分类进入简报展示。"},
            {"source": "自由时报·军武（演示）", "title": "区域安全观察聚焦防卫韧性建设", "summary": "单一来源发生访问异常时会被记录，但不会中断其他来源的采集流程。"},
        ],
    }


def international_examples() -> list[dict]:
    return [
        {
            "topic": "台湾直接相关",
            "source": "路透社 Reuters（演示）",
            "title": "Taiwan expands civil-defense readiness program",
            "english": "The program focuses on coordination among public agencies and local communities.",
            "chinese": "该计划聚焦公共部门与地方社区之间的协同准备。",
        },
        {
            "topic": "两岸议题",
            "source": "金融时报 Financial Times（演示）",
            "title": "Cross-Strait dialogue remains a focus for regional observers",
            "english": "Regional observers are tracking policy signals and their effects on stability in the Taiwan Strait.",
            "chinese": "区域观察者正追踪政策讯号及其对台海稳定的影响。",
        },
        {
            "topic": "台美关系",
            "source": "路透社 Reuters（演示）",
            "title": "United States and Taiwan discuss supply-chain cooperation",
            "english": "The talks cover resilient supply chains and practical coordination mechanisms.",
            "chinese": "会谈涵盖供应链韧性与务实协调机制。",
        },
        {
            "topic": "中美关系",
            "source": "彭博社 Bloomberg（演示）",
            "title": "China and the United States resume trade consultations",
            "english": "The consultations address trade concerns and the broader policy environment.",
            "chinese": "磋商涉及贸易关切与更广泛的政策环境。",
        },
        {
            "topic": "台湾与其他国家关系",
            "source": "华尔街日报 Wall Street Journal（演示）",
            "title": "Taiwan and Japan expand disaster-response cooperation",
            "english": "The cooperation emphasizes information sharing and emergency preparedness.",
            "chinese": "合作重点为资讯共享与紧急应变准备。",
        },
    ]


def set_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    """Mark a source/assurance table header so Word repeats it after a page break."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    """Keep each source-table row intact when Word paginates a long table."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" 页")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    run = paragraph.add_run(text)
    set_font(run, 16 if level == 1 else 13, BLUE if level == 1 else DARK_BLUE, bold=True)


def add_body(doc: Document, text: str, *, color: str = "222222", italic: bool = False, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    set_font(paragraph.add_run(text), 10.5, color, italic=italic)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_GOLD)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(title + "："), 10.5, DARK_BLUE, bold=True)
    set_font(paragraph.add_run(body), 10.5, "222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("Caption",):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string(MUTED)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("台湾新闻监测系统功能演示"), 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("演示样例，非实时监测结论  |  "), 8, MUTED)
    add_page_number(footer)


def add_cover(doc: Document, source_count: int) -> None:
    for _ in range(7):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("SYSTEM DEMONSTRATION"), 10, BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("台湾新闻监测系统\n功能演示简报"), 26, NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_font(subtitle.add_run("覆盖多源采集、完整句摘要、军武监测与国际新闻双语展示"), 12, MUTED)
    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2700, 6660])
    for index, (label, value) in enumerate(
        [("文档用途", "面向业务用户的系统能力演示"), ("来源覆盖", f"当前配置共 {source_count} 个信息来源"), ("演示边界", "不采集、不写库、不发送通知")]
    ):
        shade_cell(meta.cell(index, 0), LIGHT_BLUE)
        set_font(meta.cell(index, 0).paragraphs[0].add_run(label), 10, DARK_BLUE, bold=True)
        set_font(meta.cell(index, 1).paragraphs[0].add_run(value), 10, "222222")
    doc.add_paragraph()
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(notice.add_run(DEMO_NOTICE), 10.5, "7A5A00", bold=True)
    doc.add_page_break()


def add_flow(doc: Document) -> None:
    add_heading(doc, "一、系统全景")
    add_body(doc, "系统以来源配置为入口，对新闻逐源采集、统一去重与过滤，并形成可阅读、可追溯的业务简报。")
    flow = doc.add_table(rows=1, cols=5)
    set_table_geometry(flow, [1872, 1872, 1872, 1872, 1872])
    labels = ["多源采集", "去重与过滤", "完整句摘要", "国际相关性筛选", "Word / 通知输出"]
    for cell, label in zip(flow.rows[0].cells, labels):
        shade_cell(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.add_run(label), 10, DARK_BLUE, bold=True)
    add_callout(doc, "演示说明", "采集器单来源最多返回 20 条；单一来源异常会单独记录，不会阻断其他来源。")


def add_sources(doc: Document, rows: list[dict]) -> None:
    add_heading(doc, "二、信息来源矩阵")
    add_body(doc, "下表自动读取当前来源配置。每一行代表一个配置来源，启用状态与采集方式均按配置展示。")
    groups = ["官方来源", "台湾媒体", "军武来源", "区域中文媒体", "国际媒体", "授权通讯源"]
    for group in groups:
        entries = [row for row in rows if row["group"] == group]
        if not entries:
            continue
        if group == "授权通讯源":
            # The two newsletter rows are taller because their authorization
            # status wraps.  Start this final group on a fresh page so its
            # heading and repeated table header are never orphaned.
            doc.add_page_break()
        add_heading(doc, group, level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        set_table_geometry(table, [2350, 2150, 2900, 1960])
        headers = ["来源", "监测栏目", "采集方式", "状态"]
        for cell, text in zip(table.rows[0].cells, headers):
            shade_cell(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.add_run(text), 9.5, DARK_BLUE, bold=True)
        repeat_table_header(table.rows[0])
        for row in entries:
            cells = table.add_row().cells
            values = [row["name"], row["categories"], row["collector"], row["status"]]
            for cell, value in zip(cells, values):
                set_font(cell.paragraphs[0].add_run(value), 9, "222222")
            prevent_row_split(table.rows[-1])
        caption = doc.add_paragraph("配置来源：config/sources.yaml")
        caption.style = "Caption"


def add_category_examples(doc: Document) -> None:
    add_heading(doc, "三、分类新闻摘要示例")
    add_body(doc, "以下均为展示摘要机制的模拟新闻，示例摘要以完整句结束，避免在字数边界出现半句话。")
    for category, articles in demo_articles().items():
        add_heading(doc, category, level=2)
        for index, article in enumerate(articles, 1):
            title = doc.add_paragraph()
            title.paragraph_format.space_after = Pt(2)
            set_font(title.add_run(f"{index}. {article['title']}"), 11, NAVY, bold=True)
            add_body(doc, f"来源：{article['source']}", color=MUTED, after=2)
            add_body(doc, f"梗概：{article['summary']}", after=2)
            link = doc.add_paragraph()
            set_font(link.add_run("原文链接：来源原文链接（演示）"), 9, BLUE, italic=True)


def add_international_examples(doc: Document) -> None:
    add_heading(doc, "四、国际媒体双语摘要示例")
    add_body(doc, "国际层不会把无关海外新闻直接写入简报。系统依据台湾、两岸、台美、中美与台湾对外关系规则筛选，并仅使用已取得的英文标题与导语生成中文摘要。")
    for example in international_examples():
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [TABLE_WIDTH_DXA])
        prevent_row_split(table.rows[0])
        cell = table.cell(0, 0)
        shade_cell(cell, "F8FAFC")
        heading = cell.paragraphs[0]
        set_font(heading.add_run(example["topic"] + "｜" + example["source"]), 10.5, DARK_BLUE, bold=True)
        for label, text, color in [
            ("英文原题", example["title"], NAVY),
            ("英文摘要", example["english"], "333333"),
            ("中文摘要", example["chinese"], "333333"),
        ]:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            set_font(paragraph.add_run(label + "："), 10, DARK_BLUE, bold=True)
            set_font(paragraph.add_run(text), 10, color)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_callout(doc, "正文访问边界", "国际双语摘要只依据已采集标题和合法导语；受限媒体正文不被抓取，翻译不可用时保留英文信息并说明限制。")


def add_assurance_and_outputs(doc: Document) -> None:
    add_heading(doc, "五、运行保障与输出方式")
    safeguards = [
        ("来源故障隔离", "单来源请求设置超时与 User-Agent；失败只记录日志，不中断其他来源。"),
        ("双重去重", "先规范 URL，再以来源、日期和标题指纹识别别名链接，减少重复入库。"),
        ("内容与时效控制", "统一过滤无关社会资讯，并按新鲜、补发、过期等状态和重要度规则整理。"),
        ("摘要质量", "RSS 导语、模型摘要和元数据摘要均通过完整句规则处理，避免半句截断。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2650, 6710])
    for cell, text in zip(table.rows[0].cells, ["能力", "演示说明"]):
        shade_cell(cell, LIGHT_GRAY)
        set_font(cell.paragraphs[0].add_run(text), 10, DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for label, detail in safeguards:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(label), 10, DARK_BLUE, bold=True)
        set_font(cells[1].paragraphs[0].add_run(detail), 10, "222222")
        prevent_row_split(table.rows[-1])

    add_heading(doc, "输出与使用", level=2)
    outputs = doc.add_table(rows=1, cols=3)
    set_table_geometry(outputs, [3120, 3120, 3120])
    for cell, label, detail in zip(
        outputs.rows[0].cells,
        ["控制台摘要", "Word 简报", "飞书 / Telegram"],
        ["用于本地运行观察。", "用于分类阅读、归档与人工审阅。", "按配置发送文本、文档或重点提示。"],
    ):
        shade_cell(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.add_run(label + "\n"), 10.5, DARK_BLUE, bold=True)
        set_font(paragraph.add_run(detail), 9.5, "222222")
    add_callout(doc, "本次演示", "仅生成本地 Word 文档，不运行采集、不写入数据库，也不发送飞书、Telegram 或其他通知。")


def build_demo_docx(output_path: Path, source_rows: list[dict]) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, len(source_rows))
    add_flow(doc)
    add_sources(doc, source_rows)
    add_category_examples(doc)
    add_international_examples(doc)
    add_assurance_and_outputs(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def verify_docx(path: Path, expected_source_count: int) -> None:
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    combined = text + "\n" + table_text
    required = [
        DEMO_NOTICE,
        "信息来源矩阵",
        "分类新闻摘要示例",
        "国际媒体双语摘要示例",
        "英文原题：",
        "英文摘要：",
        "中文摘要：",
        "台湾直接相关",
        "两岸议题",
        "台美关系",
        "中美关系",
        "台湾与其他国家关系",
        "军武频道作为独立来源采集",
    ]
    missing = [item for item in required if item not in combined]
    if missing:
        raise SystemExit("文档结构核验失败，缺少：" + "、".join(missing))
    source_names = [row["name"] for row in load_source_rows()]
    absent = [name for name in source_names if name not in combined]
    if absent:
        raise SystemExit("文档未覆盖来源：" + "、".join(absent))
    print(f"结构核验通过：{expected_source_count} 个配置来源、{len(international_examples())} 类国际双语样例。")


def check_content(rows: list[dict]) -> None:
    examples = demo_articles()
    international = international_examples()
    all_summaries = [article["summary"] for articles in examples.values() for article in articles]
    all_summaries.extend(item["chinese"] for item in international)
    if any(not summary.endswith(("。", "！", "？", ".", "!", "?")) for summary in all_summaries):
        raise SystemExit("内容自检失败：存在非完整句摘要。")
    counts = Counter(row["status"] for row in rows)
    print(f"内容自检通过：来源 {len(rows)} 个；" + "；".join(f"{name} {count} 个" for name, count in sorted(counts.items())))
    print("分类样例：" + "、".join(f"{category} {len(items)} 条" for category, items in demo_articles().items()))
    print(f"国际双语议题：{len(international)} 类。")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, help="Word 输出路径")
    parser.add_argument("--check-content", action="store_true", help="只检查演示内容")
    parser.add_argument("--verify", type=Path, help="核验既有 Word 文件")
    args = parser.parse_args()
    rows = load_source_rows()
    if args.check_content:
        check_content(rows)
        return
    if args.verify:
        verify_docx(args.verify, len(rows))
        return
    output = args.output or REPORTS_DIR / f"台湾新闻监测系统功能演示_{date.today():%Y-%m-%d}.docx"
    check_content(rows)
    path = build_demo_docx(output, rows)
    print(f"Word 演示文档已生成：{path}")
    print("本次仅生成演示文档：未采集、未写库、未发送通知。")


if __name__ == "__main__":
    main()
