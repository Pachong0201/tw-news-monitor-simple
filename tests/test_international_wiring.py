"""国际媒体免费监测层 Phase I — digest / Word / main 接线测试（纯 fixture，禁止线上）。

覆盖（任务书验收项）：
- prepare_international_delivery：过滤（included/excluded）+ 跨媒体去重 + 容错回退；
- 文本 digest：不相关国际新闻不出现、canonical+coverage 标注、非国际文章不受影响；
- Word 简报：国际媒体栏目、英文标题原样、中文来源名、coverage 标注、
  国内“（四）国际新闻”与“国际媒体”无重复、原栏目（军武/宗教等）未破坏；
- main 链路语义：excluded 国际文章入库但不进 digest/Word、canonical 进 Word、
  单国际源失败不中断（复用 collect_all 现有逐源容错）。
"""

import pytest
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document

from app.database import Database
from app.digest import build_digest
from app.international import load_international_config
from app.international_events import EventCluster
from app.importance import ImportanceResult
from app.main import (
    _PrecomputedTranslationLookup,
    _build_notification_translation_lookup,
    _deliver_event_candidates,
    _notification_dedup_store,
    _precompute_international_translations,
    _translation_articles_for_delivery,
    collect_all,
    prepare_international_delivery,
)
from app.models import Article
from app.word_digest import build_word_digest
from app.notification_candidates import build_notification_candidates

PROJECT_ROOT = Path(__file__).resolve().parent.parent
CONFIG_PATH = PROJECT_ROOT / "config" / "international_media.yaml"

UTC = timezone.utc
BASE = datetime(2026, 8, 13, 10, 0, tzinfo=UTC)


def make_article(
    title: str,
    source: str,
    url: str,
    category: str = "international",
    published_at: datetime | None = None,
    summary: str | None = None,
    position: int = 1,
) -> Article:
    return Article(
        source_id="test",
        source_name=source,
        category=category,
        title=title,
        url=url,
        published_at=published_at or BASE,
        fetched_at=BASE,
        position=position,
        summary=summary,
    )


@pytest.fixture(scope="module")
def cfg():
    return load_international_config(CONFIG_PATH)


# ────────────────────────────────────────────────────────────────
# prepare_international_delivery（main.py 接线核心函数）
# ────────────────────────────────────────────────────────────────

class TestPrepareInternationalDelivery:
    def test_excludes_irrelevant_and_dedupes_same_event(self, cfg):
        arts = [
            make_article("China launches drills near Taiwan", "Reuters", "u1",
                         published_at=BASE),
            make_article("China begins military exercises around Taiwan",
                         "Financial Times", "u2",
                         published_at=BASE + timedelta(hours=1)),
            make_article("Apple stock rises on strong earnings", "Reuters", "u3"),
            make_article("中央社 本地新聞", "中央社", "d1", category="politics"),
        ]
        digest_articles, coverage = prepare_international_delivery(arts, cfg)
        # 不相关国际文章（u3）被排除；u1/u2 同一事件合并为 canonical u1；国内文章原样
        assert [a.url for a in digest_articles] == ["u1", "d1"]
        assert len(coverage["u1"]) == 2
        assert "u3" not in coverage
        assert coverage["d1"] == [arts[3]]

    def test_disabled_config_passthrough(self):
        from app.international import DISABLED_CONFIG
        arts = [make_article("Apple stock rises", "Reuters", "u1")]
        digest_articles, coverage = prepare_international_delivery(arts, DISABLED_CONFIG)
        assert digest_articles == arts
        assert coverage["u1"] == [arts[0]]

    def test_layer_exception_falls_back_to_raw(self, cfg, monkeypatch):
        """国际层内部异常 -> 回退为原样放行，绝不中断主流程。"""
        import app.main as main_module

        def _boom(articles, config):
            raise RuntimeError("layer boom")

        monkeypatch.setattr(main_module, "filter_international", _boom)
        arts = [make_article("China launches drills near Taiwan", "Reuters", "u1")]
        digest_articles, coverage = prepare_international_delivery(arts, cfg)
        assert digest_articles == arts
        assert coverage == {}


# ────────────────────────────────────────────────────────────────
# 文本 digest
# ────────────────────────────────────────────────────────────────

class TestDigestInternational:
    def test_canonical_annotated_with_coverage(self, cfg):
        u1 = make_article("China launches drills near Taiwan", "Reuters", "u1")
        u2 = make_article("China begins military exercises around Taiwan",
                          "Financial Times", "u2")
        digest = build_digest(
            [u1], BASE,
            international_coverage={u1.url: [u1, u2]},
            international_config=cfg,
        )
        assert "另有 1 家国际媒体报道同一事件" in digest
        assert "China launches drills near Taiwan" in digest

    def test_no_annotation_without_coverage(self, cfg):
        u1 = make_article("China launches drills near Taiwan", "Reuters", "u1")
        digest = build_digest(
            [u1], BASE,
            international_coverage={},
            international_config=cfg,
        )
        assert "另有" not in digest

    def test_non_international_never_annotated(self, cfg):
        d1 = make_article("中央社 本地新聞", "中央社", "d1", category="politics")
        digest = build_digest(
            [d1], BASE,
            international_coverage={d1.url: [d1]},
            international_config=cfg,
        )
        assert "另有" not in digest
        assert "中央社 本地新聞" in digest

    def test_legacy_two_arg_call_unchanged(self):
        u1 = make_article("China launches drills near Taiwan", "Reuters", "u1")
        digest = build_digest([u1], BASE)
        assert "另有" not in digest
        assert "China launches drills near Taiwan" in digest

    def test_delivery_digest_can_exclude_international_media(self, cfg):
        intl = make_article("China launches drills near Taiwan", "Reuters", "u1")
        domestic = make_article("立法院審查年度預算案", "中央社", "d1", category="politics")
        digest = build_digest(
            [intl, domestic], BASE,
            international_config=cfg,
            include_international_media=False,
        )
        assert "China launches drills near Taiwan" not in digest
        assert "立法院審查年度預算案" in digest


# ────────────────────────────────────────────────────────────────
# Word 简报
# ────────────────────────────────────────────────────────────────

class TestWordInternational:
    def _fixture_articles(self):
        return [
            make_article("China launches drills near Taiwan", "Reuters", "u1",
                         published_at=BASE,
                         summary="解放军在台海周边举行演习。"),
            make_article("China unveils new chip export controls targeting Taiwan",
                         "Financial Times", "u4",
                         published_at=BASE + timedelta(hours=2),
                         summary=None),
            make_article("立法院审查年度预算案", "中央社", "d1",
                         category="politics",
                         published_at=BASE - timedelta(hours=1)),
            make_article("美国大选初选观察：摇摆州民调胶着", "联合报", "d2",
                         category="international",
                         published_at=BASE - timedelta(hours=2)),
            make_article("国军举行年度联合演训", "自由时报·军武", "d3",
                         category="military",
                         published_at=BASE - timedelta(hours=3)),
            make_article("央行公布利率决议", "经济日报", "d5",
                         category="economy",
                         published_at=BASE - timedelta(hours=4)),
            make_article("宗教团体举办祈福法会", "中央社", "d4",
                         category="religion",
                         published_at=BASE - timedelta(hours=5)),
        ]

    def test_international_section_routing_and_rendering(self, tmp_path, cfg):
        r1, r4, d1, d2, d3, d5, d4 = self._fixture_articles()
        ft_cov = make_article("China begins military exercises around Taiwan",
                              "Financial Times", "u2",
                              published_at=BASE - timedelta(hours=1))
        bloom_cov = make_article("China starts new Taiwan military drills",
                                 "Bloomberg", "u3",
                                 published_at=BASE + timedelta(hours=1))
        coverage = {r1.url: [r1, ft_cov, bloom_cov]}

        output = build_word_digest(
            [r1, r4, d1, d2, d3, d5, d4], tmp_path,
            generated_at=BASE,
            international_config=cfg,
            international_coverage=coverage,
        )
        texts = [p.text for p in Document(output).paragraphs]

        def _has(sub):  # docx 段落文本含 “1. ” 等索引前缀，需子串匹配
            return any(sub in t for t in texts)

        # 国际媒体栏目（全部分类小节在场 -> 编号接续为（六））
        assert "（六）国际媒体" in texts
        # 英文标题原样输出，不做翻译
        assert _has("China launches drills near Taiwan")
        assert _has("China unveils new chip export controls targeting Taiwan")
        # 中文来源名 + 英文 canonical 名
        assert _has("来源：路透社（Reuters）")
        assert _has("来源：金融时报（Financial Times）")
        # coverage 标注（用 display_name）
        assert _has("另据金融时报、彭博社等报道同一事件")
        # 双语摘要：英文原摘要始终展示；没有真实翻译器时不伪造中文摘要。
        assert _has("英文摘要：解放军在台海周边举行演习。")
        assert not any("英文摘要：None" in t for t in texts)
        assert not any("中文摘要：" in t for t in texts)
        # 栏目内排序 (published_at desc)：12:00 的 r4 在 10:00 的 r1 之前
        idx_r4 = next(i for i, t in enumerate(texts) if "China unveils new chip export controls targeting Taiwan" in t)
        idx_r1 = next(i for i, t in enumerate(texts) if "China launches drills near Taiwan" in t)
        assert idx_r4 < idx_r1
        # 原栏目未破坏
        assert "（一）政治新闻" in texts
        assert "（二）经济新闻" in texts
        assert "（三）军武" in texts
        assert "（四）国际新闻" in texts
        assert "（五）宗教" in texts
        # 国内“（四）国际新闻”与“国际媒体”无重复：
        # 国际媒体文章只出现一次且位于国际媒体栏目内；（四）内只有国内国际新闻
        assert sum("China launches drills near Taiwan" in t for t in texts) == 1
        assert sum("美国大选初选观察：摇摆州民调胶着" in t for t in texts) == 1
        intl_cat_idx = texts.index("（四）国际新闻")
        rel_idx = texts.index("（五）宗教")
        between = texts[intl_cat_idx:rel_idx]
        assert not any("China" in t for t in between)
        intl_sec_idx = texts.index("（六）国际媒体")
        r1_idx = next(i for i, t in enumerate(texts) if "China launches drills near Taiwan" in t)
        assert intl_sec_idx < r1_idx

    def test_no_config_keeps_legacy_category_routing(self, tmp_path):
        """未传 international_config：国际媒体文章仍按旧行为进分类小节。"""
        r1 = make_article("China launches drills near Taiwan", "Reuters", "u1")
        d2 = make_article("美国大选初选观察", "联合报", "d2", category="international")
        output = build_word_digest([r1, d2], tmp_path, generated_at=BASE)
        texts = [p.text for p in Document(output).paragraphs]
        assert "（一）国际新闻" in texts
        assert any("China launches drills near Taiwan" in t for t in texts)
        assert any("来源：Reuters" in t for t in texts)  # 旧样式：直接英文来源名
        assert not any("国际媒体" in t and "来源" not in t for t in texts)


# ────────────────────────────────────────────────────────────────
# main 接线语义：落库 / digest / Word / 单源失败隔离
# ────────────────────────────────────────────────────────────────

class TestMainWiring:
    def test_missing_translator_precomputes_strict_metadata_fallback(self, cfg):
        article = make_article("Taiwan update", "Reuters", "u-fallback")
        article.summary = "A public teaser about Taiwan."
        translations = _precompute_international_translations([article], cfg, translator=None)
        assert translations[article.url].status == "fallback"
        assert translations[article.url].cn_title == article.title
        assert translations[article.url].cn_summary == article.summary

    def test_old_canonical_and_fresh_coverage_share_precomputed_translation_policy(self, cfg):
        old = make_article(
            "US approves arms sales to Taiwan", "Reuters", "old-canonical",
            published_at=BASE - timedelta(hours=2),
        )
        fresh = make_article(
            "US approves new arms sales package for Taiwan", "Financial Times", "fresh-coverage",
            published_at=BASE,
        )
        coverage = {old.url: [old, fresh]}
        members = _translation_articles_for_delivery([old], coverage)
        translations = _precompute_international_translations(members, cfg, translator=None)
        assert set(translations) == {old.url, fresh.url}
        assert translations[old.url].status == "fallback"
        assert translations[fresh.url].status == "fallback"
        # Notification lookup deliberately exposes the same fallback policy;
        # it cannot silently invent a second Chinese result for fresh evidence.
        lookup = _build_notification_translation_lookup(members, translations, coverage)
        assert lookup is not None
        assert lookup.translate(
            fresh.title, fresh.summary, source_name=fresh.source_name
        )[0] == old.title

        cluster = EventCluster(
            event_id="evt-old-fresh",
            canonical=old,
            members=[old, fresh],
            coverage=[old, fresh],
        )
        candidates = build_notification_candidates(
            [cluster],
            [
                (old, ImportanceResult(score=90, level="critical", reasons=["arms"])),
                (fresh, ImportanceResult(score=90, level="critical", reasons=["arms"])),
            ],
            {"fresh_articles": [fresh], "catch_up_urls": {old.url}},
            BASE,
            translator=lookup,
        )
        assert len(candidates) == 1
        assert candidates[0].cn_title == old.title

    def test_notification_dedup_path_is_isolated_for_dry_run(self, tmp_path, cfg):
        assert _notification_dedup_store(tmp_path, cfg, dry_run=True) is None
        store = _notification_dedup_store(tmp_path, cfg, dry_run=False)
        assert store is not None
        assert store._path == tmp_path / "data" / "international_notification_dedup.json"

    def test_event_dedup_marks_only_explicit_true_success(self):
        class Candidate:
            dedup_key = "event:test"

        class Store:
            def __init__(self):
                self.marked = []
            def mark_sent(self, key, now=None):
                self.marked.append(key)

        class Notifier:
            def __init__(self, result=None, error=False):
                self.result = result
                self.error = error
            def send_event_candidates(self, candidates):
                if self.error:
                    raise RuntimeError("send failed")
                return self.result

        now = BASE
        for result, error in ((None, False), (False, False), (None, True), (True, False)):
            store = Store()
            marked = _deliver_event_candidates(
                Notifier(result, error), [Candidate()], store, now
            )
            assert marked is (result is True and not error)
            assert store.marked == (["event:test"] if result is True and not error else [])

    def test_excluded_intl_article_saved_to_db_but_not_in_digest_or_word(
        self, tmp_path, cfg
    ):
        arts = [
            make_article("China launches drills near Taiwan", "Reuters", "u1",
                         published_at=BASE),
            make_article("China begins military exercises around Taiwan",
                         "Financial Times", "u2",
                         published_at=BASE + timedelta(hours=1)),
            make_article("Apple stock rises on strong earnings", "Reuters", "u3"),
            make_article("中央社 本地新聞", "中央社", "d1", category="politics"),
        ]
        # 主流程语义：所有国际文章（含不相关的）按 URL 去重后正常入库
        db = Database(tmp_path / "wiring.db")
        db.connect()
        db.create_tables()
        try:
            inserted = db.save_articles(arts)
            assert len(inserted) == 4
            row = db.conn.execute(
                "SELECT COUNT(*) FROM articles WHERE url = 'u3'"
            ).fetchone()
            assert row[0] == 1
        finally:
            db.close()

        # 国际层：u3 被排除；u1/u2 合并 -> canonical u1
        digest_articles, coverage = prepare_international_delivery(arts, cfg)
        assert [a.url for a in digest_articles] == ["u1", "d1"]

        # 文本 digest：excluded 不出现，canonical 标注 coverage，非国际文章原样
        digest = build_digest(
            digest_articles, BASE,
            international_coverage=coverage,
            international_config=cfg,
        )
        assert "另有 1 家国际媒体报道同一事件" in digest
        assert "China launches drills near Taiwan" in digest
        assert "Apple stock rises" not in digest
        assert "中央社 本地新聞" in digest

        # Word：canonical 进简报，excluded 不出现
        output = build_word_digest(
            digest_articles, tmp_path / "out",
            generated_at=BASE,
            international_config=cfg,
            international_coverage=coverage,
        )
        texts = [p.text for p in Document(output).paragraphs]
        assert any("China launches drills near Taiwan" in t for t in texts)
        assert not any("Apple stock rises" in t for t in texts)
        assert any("中央社 本地新聞" in t for t in texts)

    def test_single_international_source_failure_does_not_break_run(
        self, tmp_path, monkeypatch, cfg
    ):
        """单国际源失败不中断（复用 collect_all 现有逐源容错语义）。"""
        class FailingIntlCollector:
            def __init__(self, _source):
                pass

            def collect(self):
                raise RuntimeError("reuters feed down")

            def close(self):
                pass

        class GoodDomesticCollector:
            def __init__(self, source_cfg):
                self.source = source_cfg

            def collect(self):
                now = datetime(2026, 8, 13, 10, 0, tzinfo=UTC)
                return [
                    make_article("China launches drills near Taiwan", "Reuters",
                                 "https://example.com/u1", published_at=now),
                    make_article("立法院审查年度预算案", "中央社",
                                 "https://example.com/d1", category="politics",
                                 published_at=now),
                ]

            def close(self):
                pass

        monkeypatch.setattr(
            "app.main.COLLECTOR_MAP",
            {"reuters": FailingIntlCollector, "rss": GoodDomesticCollector},
        )
        db = Database(tmp_path / "intl_fail.db")
        db.connect()
        db.create_tables()
        try:
            result = collect_all(
                [
                    {"id": "reuters_international", "name": "Reuters",
                     "type": "reuters", "url": "https://www.reuters.com/feed",
                     "enabled": True, "category": "international"},
                    {"id": "cna_politics", "name": "中央社",
                     "type": "rss", "url": "https://example.com/cna.xml",
                     "enabled": True, "category": "politics"},
                ],
                db,
            )
            inserted, total, _dup, failed = result[:4]
            assert failed == ["reuters_international"]
            assert total == 2
            assert len(inserted) == 2
            assert db.count_articles() == 2
        finally:
            db.close()

        # 主流程后续（国际层）照常运行，采集到的国际文章正常进简报
        digest_articles, _coverage = prepare_international_delivery(inserted, cfg)
        assert [a.url for a in digest_articles] == [
            "https://example.com/u1", "https://example.com/d1",
        ]
