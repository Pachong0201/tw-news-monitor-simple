import pytest
import tempfile
from datetime import datetime, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo

from app.freshness import FreshnessResult, filter_fresh_articles
from app.models import Article
from app.word_digest import build_word_digest

TAIPEI = ZoneInfo("Asia/Taipei")
RUN = datetime(2026, 7, 19, 10, 0, tzinfo=TAIPEI)

def make(pub, url="https://test.com/a", title="T", src="test", sid="test", pos=1):
    pub_aware = pub.replace(tzinfo=TAIPEI) if pub and pub.tzinfo is None else pub
    return Article(
        source_id=sid, source_name=src, category="politics",
        title=title, url=url, published_at=pub_aware,
        fetched_at=RUN, position=pos,
    )

class TestCatchUpClassification:
    def test_60min_fresh(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=60))], RUN, catch_up_enabled=True)
        assert len(fr.fresh_articles) == 1
        assert len(fr.catch_up_articles) == 0
        assert len(fr.stale_articles) == 0

    def test_120min_catch_up(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=120))], RUN, catch_up_enabled=True)
        assert len(fr.fresh_articles) == 0
        assert len(fr.catch_up_articles) == 1
        assert len(fr.stale_articles) == 0

    def test_600min_catch_up(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=600))], RUN, catch_up_enabled=True)
        assert len(fr.fresh_articles) == 0
        assert len(fr.catch_up_articles) == 1
        assert len(fr.stale_articles) == 0

    def test_800min_stale(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=800))], RUN, catch_up_enabled=True)
        assert len(fr.fresh_articles) == 0
        assert len(fr.catch_up_articles) == 0
        assert len(fr.stale_articles) == 1

    def test_91min_is_catch_up(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=91))], RUN, catch_up_enabled=True)
        assert len(fr.catch_up_articles) == 1

    def test_720min_exact_catch_up(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=720))], RUN, catch_up_enabled=True, catch_up_max_minutes=720)
        assert len(fr.catch_up_articles) == 1

    def test_721min_stale(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=721))], RUN, catch_up_enabled=True)
        assert len(fr.catch_up_articles) == 0
        assert len(fr.stale_articles) == 1

    def test_no_pub_unknown(self):
        a = make(RUN - timedelta(minutes=120))
        a.published_at = None
        fr = filter_fresh_articles([a], RUN, catch_up_enabled=True)
        assert len(fr.unknown_time_articles) == 1

    def test_future_abnormal(self):
        fr = filter_fresh_articles([make(RUN + timedelta(minutes=20))], RUN, catch_up_enabled=True)
        assert len(fr.future_time_articles) == 1

    def test_catch_up_disabled_stale(self):
        fr = filter_fresh_articles([make(RUN - timedelta(minutes=120))], RUN, catch_up_enabled=False)
        assert len(fr.catch_up_articles) == 0
        assert len(fr.stale_articles) == 1

    def test_invalid_config_raises(self):
        with pytest.raises(ValueError, match="must be greater than"):
            filter_fresh_articles([make(RUN - timedelta(minutes=60))], RUN, catch_up_enabled=True, catch_up_max_minutes=60)

    def test_mixed_categories(self):
        arts = [
            make(RUN - timedelta(minutes=30), url="a", title="fresh"),
            make(RUN - timedelta(minutes=120), url="b", title="catch"),
            make(RUN - timedelta(minutes=800), url="c", title="old"),
        ]
        fr = filter_fresh_articles(arts, RUN, catch_up_enabled=True)
        assert len(fr.fresh_articles) == 1
        assert len(fr.catch_up_articles) == 1
        assert len(fr.stale_articles) == 1

class TestCatchUpWord:
    def test_catch_up_prefix(self):
        arts = [make(RUN - timedelta(minutes=120), url="https://t.com/1", title="bw1")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls={"https://t.com/1"})
            from docx import Document
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any(("【补发】" in t) for t in texts)

    def test_fresh_no_prefix(self):
        arts = [make(RUN - timedelta(minutes=30), url="https://t.com/2", title="fresh2")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls=set())
            from docx import Document
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert "1. fresh2" in " ".join(texts)

    def test_catch_up_status(self):
        arts = [make(RUN - timedelta(minutes=120), url="https://t.com/3", title="bw3")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls={"https://t.com/3"})
            from docx import Document
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any(("状态：补发" in t) for t in texts)

    def test_stats_show_catch_up(self):
        arts = [
            make(RUN - timedelta(minutes=30), url="https://t.com/f1", title="f1"),
            make(RUN - timedelta(minutes=120), url="https://t.com/c1", title="c1"),
        ]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls={"https://t.com/c1"})
            from docx import Document
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            combined = " ".join(texts)
            assert "正常新闻：1条" in combined
            assert "补发新闻：1条" in combined

    def test_mixed_sort(self):
        arts = [
            make(RUN - timedelta(minutes=30), url="a", title="f30"),
            make(RUN - timedelta(minutes=120), url="b", title="c120"),
            make(RUN - timedelta(minutes=15), url="c", title="f15"),
        ]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls={"b"})
            from docx import Document
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            combined = " ".join(texts)
            i15 = combined.index("f15")
            i30 = combined.index("f30")
            i120 = combined.index("c120")
            assert i15 < i30 < i120

    def test_only_catch_up_generates_word(self):
        arts = [make(RUN - timedelta(minutes=120), url="https://t.com/only")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(arts, Path(tmp), generated_at=RUN, catch_up_urls={"https://t.com/only"})
            assert output.exists()
            assert output.suffix == ".docx"

    def test_no_articles_raises(self):
        with tempfile.TemporaryDirectory() as tmp:
            with pytest.raises(ValueError, match="No articles"):
                build_word_digest([], Path(tmp))
