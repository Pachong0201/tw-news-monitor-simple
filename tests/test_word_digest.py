import tempfile
from datetime import datetime
from hashlib import sha256
from pathlib import Path

import pytest
from docx import Document

from app.models import Article
from app.word_digest import build_word_digest
from app.international import load_international_config
from app.international_translation import FakeTranslator, TranslationResult, translate_article


def make_article(**kwargs):
    now = datetime(2026, 7, 14, 20, 30)
    return Article(
        source_id=kwargs.get("source_id", "test"),
        source_name=kwargs.get("source_name", "中央社"),
        category=kwargs.get("category", "politics"),
        title=kwargs.get("title", "測試新聞"),
        url=kwargs.get("url", "https://example.com/news/1"),
        published_at=kwargs.get("published_at", now),
        fetched_at=now,
        position=kwargs.get("position", 1),
    )


class TestWordDigest:
    def test_generates_docx(self):
        """1. Successfully generate a .docx file."""
        articles = [make_article()]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            assert output.suffix == ".docx"
            assert output.exists()

    def test_auto_creates_output_dir(self):
        """2. Output directory is created if not exists."""
        articles = [make_article()]
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp) / "reports" / "sub"
            output = build_word_digest(articles, output_dir)
            assert output.parent.exists()

    def test_contains_title(self):
        """3. Word contains the main title."""
        articles = [make_article()]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("台湾新闻监测" in t for t in texts)

    def test_contains_categories(self):
        """4. Word contains politics, economy, international."""
        articles = [
            make_article(category="politics", url="https://example.com/1"),
            make_article(category="economy", url="https://example.com/2"),
            make_article(category="international", url="https://example.com/3"),
        ]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("政治" in t for t in texts)
            assert any("经济" in t for t in texts)
            assert any("国际" in t for t in texts)

    def test_contains_article_title(self):
        """5. Word contains article titles."""
        articles = [make_article(title="獨家新聞標題")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("獨家新聞標題" in t for t in texts)

    def test_contains_media_name(self):
        """6. Word contains media name."""
        articles = [make_article(source_name="中央社")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("中央社" in t for t in texts)

    def test_contains_publish_time(self):
        """7. Word contains publish time."""
        articles = [make_article(published_at=datetime(2026, 7, 14, 20, 30))]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("2026-07-14 20:30" in t for t in texts)

    def test_contains_url(self):
        """8. Word contains original link text."""
        articles = [make_article(url="https://example.com/test-article")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("https://example.com/test-article" in t for t in texts)

    def test_chinese_text(self):
        """9. Chinese text is written correctly."""
        articles = [make_article(title="繁體中文測試新聞")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("繁體中文測試新聞" in t for t in texts)

    def test_filename_format(self):
        """10. Filename format is correct."""
        articles = [make_article()]
        dt = datetime(2026, 7, 14, 20, 30)
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp), generated_at=dt)
            assert "台湾新闻监测_2026-07-14_2030" in output.name
            assert output.name.endswith(".docx")

    def test_no_articles_raises(self):
        """13. No articles raises ValueError, no file generated."""
        with tempfile.TemporaryDirectory() as tmp:
            with pytest.raises(ValueError, match="No articles"):
                build_word_digest([], Path(tmp))

    def test_contains_footer(self):
        """Word contains the footer disclaimer."""
        articles = [make_article()]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("本简报由自动新闻监测程序生成" in t for t in texts)

    def test_count_meta(self):
        """Meta line shows correct article count."""
        articles = [make_article() for _ in range(5)]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            assert any("5条" in t for t in texts)

    def test_hyperlink_in_xml(self):
        """URL is added as a clickable hyperlink."""
        articles = [make_article(url="https://example.com/hyperlink-test")]
        with tempfile.TemporaryDirectory() as tmp:
            output = build_word_digest(articles, Path(tmp))
            doc = Document(str(output))
            xml = doc.element.xml
            assert "w:hyperlink" in xml

    def test_international_missing_translation_uses_english_metadata_only(self, tmp_path):
        """A missing mapping never invents a Chinese title or summary."""
        config = load_international_config()
        article = make_article(
            source_id="reuters_international",
            source_name="Reuters",
            category="international",
            title="China announces new measures for Taiwan",
            url="https://example.test/metadata-only",
        )
        article.summary = "The statement described measures for Taiwan."
        output = build_word_digest(
            [article], tmp_path,
            international_config=config,
            international_translations={},
        )
        texts = [p.text for p in Document(output).paragraphs]
        assert any("英文原题：China announces new measures for Taiwan" in text for text in texts)
        assert any("英文摘要：The statement described measures for Taiwan." in text for text in texts)
        assert not any("中文摘要：" in text for text in texts)
        assert not any("中国" in text and "标题" in text for text in texts)

    def test_international_real_translation_mapping_renders_chinese_fields(self, tmp_path):
        config = load_international_config()
        article = make_article(
            source_id="reuters_international",
            source_name="Reuters",
            category="international",
            title="Taiwan military drills",
            url="https://example.test/translated",
        )
        article.summary = "Military drills were announced near Taiwan."
        translated = translate_article(article, FakeTranslator())
        output = build_word_digest(
            [article], tmp_path,
            international_config=config,
            international_translations={article.url: translated},
        )
        texts = [p.text for p in Document(output).paragraphs]
        assert any("台湾 军事演习" in text for text in texts)
        assert any("英文摘要：Military drills were announced near Taiwan." in text for text in texts)
        assert any("中文摘要：" in text for text in texts)
        assert any("系统判断：" in text for text in texts)

    def test_llm_summary_is_not_mislabeled_as_english_summary(self, tmp_path):
        config = load_international_config()
        article = make_article(
            source_name="Reuters",
            title="Taiwan military drills begin",
            summary="这是既有的中文模型梗概。",
        )
        article.summary_source = "llm"
        translated = TranslationResult(
            status="translated",
            cn_title="台湾军演启动",
            cn_summary="台湾启动军事演习。",
            limitation="",
            body_fetch_count=0,
        )
        output = build_word_digest(
            [article], tmp_path, international_config=config,
            international_translations={article.url: translated},
        )
        texts = [p.text for p in Document(output).paragraphs]
        assert any("英文摘要：未提供合法英文摘要。" in text for text in texts)
        assert any("中文摘要：台湾启动军事演习。" in text for text in texts)
        assert not any("英文摘要：这是既有的中文模型梗概。" in text for text in texts)

    def test_fallback_mapping_cannot_smuggle_chinese_fields(self, tmp_path):
        config = load_international_config()
        article = make_article(
            source_id="reuters_international",
            source_name="Reuters",
            category="international",
            title="Taiwan update",
            url="https://example.test/fallback-status",
        )
        article.summary = "An English teaser."
        fallback = TranslationResult(
            cn_title="伪造中文标题",
            cn_summary="伪造中文摘要",
            status="fallback",
            limitation="provider unavailable",
            body_fetch_count=0,
        )
        output = build_word_digest(
            [article], tmp_path,
            international_config=config,
            international_translations={article.url: fallback},
        )
        texts = [p.text for p in Document(output).paragraphs]
        assert not any("伪造中文" in text for text in texts)
        assert any("英文原题：Taiwan update" in text for text in texts)
        assert any("英文摘要：An English teaser." in text for text in texts)
        assert not any("中文摘要：" in text for text in texts)

    def test_same_input_produces_byte_deterministic_docx(self, tmp_path):
        config = load_international_config()
        article = make_article(
            source_id="reuters_international",
            source_name="Reuters",
            category="international",
            title="Taiwan update",
            url="https://example.test/deterministic",
            published_at=datetime(2026, 7, 14, 20, 30),
        )
        article.summary = "A public teaser."
        generated_at = datetime(2026, 8, 15, 10, 0)
        first = build_word_digest(
            [article], tmp_path / "one", generated_at=generated_at,
            international_config=config,
        )
        second = build_word_digest(
            [article], tmp_path / "two", generated_at=generated_at,
            international_config=config,
        )
        assert sha256(first.read_bytes()).digest() == sha256(second.read_bytes()).digest()
