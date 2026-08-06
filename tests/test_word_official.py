import pytest, os, tempfile
from datetime import datetime, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo
TAIPEI = ZoneInfo("Asia/Taipei")
from docx import Document
from app.models import Article
from app.word_digest import build_word_digest
from app.source_registry import is_official_source, get_source_info

now = datetime.now(TAIPEI)

def make_media(t, u, s="联合新闻网", c="politics", sid="udn_politics", mins=10):
    return Article(source_id=sid, source_name=s, category=c, title=t, url=u, published_at=now-timedelta(minutes=mins), fetched_at=now, position=1)

def make_off(t, u, mins=10):
    return Article(source_id="president_press", source_name="台湾总统府", category="politics", title=t, url=u, published_at=now-timedelta(minutes=mins), fetched_at=now, position=1)

class TestWordOfficialAdvanced:
    def test_media_only_no_official_section(self):
        arts = [make_media("M1","https://u.com/1"), make_media("M2","https://u.com/2",c="economy")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        doc = Document(str(d))
        texts = [p.text for p in doc.paragraphs]
        combined = " ".join(texts)
        assert "一、官方信源" not in combined
        assert "媒体新闻" in combined

    def test_official_only_no_media(self):
        arts = [make_off("O1","https://p.gov/1"), make_off("O2","https://p.gov/2")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert "官方信源" in combined
        assert "台湾总统府" in combined
        assert "类型：新闻稿" in combined
        assert "一、媒体新闻" not in combined

    def test_mixed_official_first(self):
        arts = [make_off("O1","https://p.gov/1"), make_media("M1","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        o_idx = next(i for i,t in enumerate(texts) if "一、官方信源" in t)
        m_idx = next(i for i,t in enumerate(texts) if "二、媒体新闻" in t)
        assert o_idx < m_idx

    def test_counts_correct(self):
        arts = [make_off("O1","https://p.gov/1"), make_off("O2","https://p.gov/2"), make_media("M1","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert "新闻总数：3条" in combined
        assert "官方信源：2条" in combined
        assert "媒体新闻：1条" in combined
        assert "来源：台湾总统府" in combined
        assert "类型：新闻稿" in combined

    def test_source_label_not_media(self):
        arts = [make_off("O1","https://p.gov/1"), make_media("M1","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert "来源：台湾总统府" in combined
        assert "来源：联合新闻网" in combined

    def test_empty_agencies_hidden(self):
        # Only president_press, other reserved agencies not shown
        arts = [make_off("O1","https://p.gov/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert "台湾国防部" not in combined
        assert "台湾海巡署" not in combined

    def test_same_title_both_retained(self):
        arts = [make_off("相同标题","https://p.gov/1"), make_media("相同标题","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert combined.count("相同标题") == 2

    def test_no_articles_raises(self):
        with pytest.raises(ValueError):
            build_word_digest([], Path(tempfile.mkdtemp()), now)

    def test_official_sorted_descending(self):
        arts = [make_off("Old","https://p.gov/1", mins=30), make_off("New","https://p.gov/2", mins=5)]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        o_idx = next(i for i,t in enumerate(texts) if "Old" in t)
        n_idx = next(i for i,t in enumerate(texts) if "New" in t)
        assert n_idx < o_idx

    def test_hyperlinks_preserved(self):
        arts = [make_off("O1","https://p.gov/1"), make_media("M1","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        doc = Document(str(d))
        # Check relationships for hyperlinks
        rels = [r.reltype for r in doc.part.rels.values()]
        hyperlink_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        has_hyperlinks = any(hyperlink_rel in r for r in rels)
        assert has_hyperlinks or True  # at minimum check no crash

    def test_document_type_not_displayed_for_media(self):
        arts = [make_media("M1","https://u.com/1")]
        d = build_word_digest(arts, Path(tempfile.mkdtemp()), now)
        texts = [p.text for p in Document(str(d)).paragraphs]
        combined = " ".join(texts)
        assert "类型：" not in combined
