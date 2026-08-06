from pathlib import Path

import pytest
from docx import Document

from app.assessment.security_scan import scan_text
from app.assessment.word_report_renderer import (
    DRAFT_LABEL,
    SECTION_ORDER,
    extract_word_body,
    extract_word_text,
    render_word_report,
)
from tests.assessment.llm.conftest import build_contract, make_report


class TestWordReportRenderer:
    def test_draft_render_creates_docx(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        path = Path(info["docx_path"])
        assert path.exists()
        assert "数据不完整草稿" in path.name
        Document(path)

    def test_draft_label_present(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        assert DRAFT_LABEL in text

    def test_final_report_no_draft_label(self, tmp_path):
        report = make_report(build_contract(), fixture="valid_final")
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        assert "数据不完整草稿" not in Path(info["docx_path"]).name
        assert DRAFT_LABEL not in extract_word_text(Path(info["docx_path"]))

    def test_rejected_report_refused(self, tmp_path):
        report = make_report(build_contract())
        report["report_status"] = "rejected"
        with pytest.raises(ValueError):
            render_word_report(report, output_dir=tmp_path, mode="development")

    def test_title_matches(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        assert report["title"] in extract_word_text(Path(info["docx_path"]))

    def test_eight_sections_present(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        for heading in SECTION_ORDER:
            assert heading in text
        assert info["section_count"] == 8

    def test_all_claims_rendered_in_section_order(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        body = extract_word_body(Path(info["docx_path"]))
        expected = []
        for section in report["sections"]:
            for cid in section["claim_ids"]:
                if cid not in expected:
                    expected.append(cid)
        positions = []
        claims = {c["claim_id"]: c for c in report["claims"]}
        for cid in expected:
            positions.append(body.find(claims[cid]["claim_text"]))
        assert positions == sorted(positions)
        assert info["rendered_claim_count"] == len(report["claims"])

    def test_required_disclosures_rendered(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        claims = {c["claim_id"]: c for c in report["claims"]}
        for cid in report["required_disclosures"]:
            assert claims[cid]["claim_text"] in text

    def test_evidence_appendix_present(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        doc = Document(Path(info["docx_path"]))
        assert len(doc.tables) >= 2  # 信息栏 + 证据附录
        assert "证据附录" in extract_word_text(Path(info["docx_path"]))

    def test_long_event_id_not_crashing(self, tmp_path):
        report = make_report(build_contract())
        long_id = "EVENT_" + "X" * 300
        report["claims"][6]["supporting_event_ids"] = [long_id]
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        assert long_id in text

    def test_chinese_not_garbled(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        assert "陈亭妃" in text or "谢龙介" in text
        assert "\ufffd" not in text

    def test_no_secrets_or_paths_in_docx_text(self, tmp_path):
        report = make_report(build_contract())
        info = render_word_report(report, output_dir=tmp_path, mode="development")
        text = extract_word_text(Path(info["docx_path"]))
        scan = scan_text(text)
        assert not any(
            scan[k]
            for k in (
                "deepseek_api_key_exposed",
                "feishu_webhook_exposed",
                "authorization_header_exposed",
                "absolute_developer_path_exposed",
            )
        )

    def test_renderer_does_not_call_model(self):
        import inspect

        assert "provider" not in inspect.signature(render_word_report).parameters
