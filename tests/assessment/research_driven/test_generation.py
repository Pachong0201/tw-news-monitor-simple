"""research_driven 生成编排端到端测试（mock 适配器，冻结正式输入只读）。"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.assessment.research_driven.generation import run_generation
from app.assessment.r2.state import ReportRunStore

PROJECT_ROOT = Path(__file__).resolve().parents[3]
CONFIG_PATH = PROJECT_ROOT / "config/election_assessment.yaml"


def _run(tmp_path: Path, **kwargs) -> dict:
    defaults = dict(
        config_path=CONFIG_PATH,
        runs_root=tmp_path,
        as_of=date(2026, 8, 9),
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 31),
        trigger_type="controlled",
        check_only=False,
        force_regenerate=False,
        project_root=PROJECT_ROOT,
        mock_fixture="valid",
        # 测试使用独立锁目录，避免与生产运行共用锁键
        lock_dir=tmp_path / "locks",
    )
    defaults.update(kwargs)
    return run_generation(**defaults)


def test_end_to_end_generation_and_artifacts(tmp_path):
    result = _run(tmp_path)
    assert result["code"] == "GENERATED_READY_FOR_REVIEW"
    assert result["fact_safety_status"] == "pass"
    period_dir = tmp_path / "periods" / "20260716_20260731"
    for name in (
        "input_manifest.json",
        "research_pack.json",
        "ASSESSMENT_RESEARCH_PACK.md",
        "analysis_plan.json",
        "final_article.md",
        "final_article.docx",
        "fact_safety_audit.json",
        "review_notes.json",
        "run_metadata.json",
    ):
        assert (period_dir / name).exists(), name
    # 生产根预览副本
    assert (tmp_path / "FINAL_ASSESSMENT_PREVIEW.md").exists()
    assert (tmp_path / "FINAL_ASSESSMENT_PREVIEW.docx").exists()
    assert (tmp_path / "ASSESSMENT_RESEARCH_PACK.md").exists()
    # run 记录
    store = ReportRunStore(tmp_path)
    run = store.get("tainan_mayoral_2026__20260716__20260731")
    assert run["generation_status"] == "ready_for_review"
    assert run["facts_cutoff"] >= "2026-07-31"
    assert run["research_pack_hash"]
    assert run["article_hash"]
    assert run["word_hash"]


def test_research_pack_markdown_contains_facts(tmp_path):
    result = _run(tmp_path)
    assert result["code"] == "GENERATED_READY_FOR_REVIEW"
    pack_md = (tmp_path / "periods/20260716_20260731/ASSESSMENT_RESEARCH_PACK.md").read_text(encoding="utf-8")
    assert "报告周期：2026-07-16 至 2026-07-31" in pack_md
    assert "facts_cutoff" in pack_md


def test_word_contains_article_and_no_internal_ids(tmp_path):
    result = _run(tmp_path)
    assert result["code"] == "GENERATED_READY_FOR_REVIEW"
    word_path = tmp_path / "periods/20260716_20260731/final_article.docx"
    doc = Document(word_path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    assert "核心判断" in text
    assert "报告周期" in text
    # Word 正文不得出现内部标识
    assert "evt_" not in text
    assert "src_" not in text
    assert "claim_id" not in text


def test_idempotent_skip_and_force_regenerate(tmp_path):
    first = _run(tmp_path)
    assert first["code"] == "GENERATED_READY_FOR_REVIEW"
    second = _run(tmp_path, trigger_type="scheduled")
    assert second["code"] == "SKIPPED_ALREADY_GENERATED"
    third = _run(tmp_path, force_regenerate=True)
    assert third["code"] == "GENERATED_READY_FOR_REVIEW"
    store = ReportRunStore(tmp_path)
    run = store.get("tainan_mayoral_2026__20260716__20260731")
    # force 后 run_id 保留（同一 run 记录续跑）
    assert run["run_id"] == first["run_id"]
    # 旧记录已快照进 history
    history = list((tmp_path / "history").glob("*.json"))
    assert len(history) >= 1


def test_period_not_ready_gate(tmp_path):
    # period_end 超出 facts_cutoff → REPORT_PERIOD_NOT_READY，不得生成文章
    result = _run(
        tmp_path,
        period_start=date(2026, 12, 1),
        period_end=date(2026, 12, 15),
    )
    assert result["code"] == "REPORT_PERIOD_NOT_READY"
    period_dir = tmp_path / "periods" / "20261201_20261215"
    assert not (period_dir / "final_article.md").exists()
    store = ReportRunStore(tmp_path)
    run = store.get("tainan_mayoral_2026__20261201__20261215")
    assert run["generation_status"] == "period_not_ready"


def test_generation_failed_preserves_research_pack(tmp_path):
    result = _run(tmp_path, mock_fixture="api_failure")
    assert result["code"] == "GENERATION_FAILED"
    assert result["research_pack_ready"] is True
    assert result["article_generation_failed"] is True
    period_dir = tmp_path / "periods" / "20260716_20260731"
    assert (period_dir / "research_pack.json").exists()
    assert (period_dir / "ASSESSMENT_RESEARCH_PACK.md").exists()
    assert not (period_dir / "final_article.md").exists()
    store = ReportRunStore(tmp_path)
    run = store.get("tainan_mayoral_2026__20260716__20260731")
    assert run["generation_status"] == "generation_failed"


def test_future_leakage_hard_block(tmp_path):
    result = _run(tmp_path, mock_fixture="future_leakage")
    assert result["code"] == "MACHINE_HARD_BLOCKED"
    assert any("future_event_leakage" in r for r in result["hard_block_reasons"])
    store = ReportRunStore(tmp_path)
    run = store.get("tainan_mayoral_2026__20260716__20260731")
    assert run["generation_status"] == "machine_rejected"
    # 研究包与审计仍保留（机器拒收不丢研究）
    period_dir = tmp_path / "periods" / "20260716_20260731"
    assert (period_dir / "research_pack.json").exists()
    assert (period_dir / "fact_safety_audit.json").exists()


def test_analysis_plan_structure(tmp_path):
    result = _run(tmp_path)
    assert result["code"] == "GENERATED_READY_FOR_REVIEW"
    plan = json.loads(
        (tmp_path / "periods/20260716_20260731/analysis_plan.json").read_text(encoding="utf-8")
    )
    assert plan["primary_thesis"]["judgment"]
    assert plan["key_changes"]
    assert plan["causal_chain"]
    assert plan["trend_outlook"]["short_term"]
    assert plan["trend_outlook"]["medium_term"]
    assert plan["trend_outlook"]["key_turning_conditions"]
    article = json.loads(
        (tmp_path / "periods/20260716_20260731/research_pack.json").read_text(encoding="utf-8")
    )
    assert article["pack_schema_version"] == "1.0"


def test_non_schedule_day_requires_explicit_period(tmp_path):
    with pytest.raises(ValueError):
        run_generation(
            config_path=CONFIG_PATH,
            runs_root=tmp_path,
            as_of=date(2026, 8, 10),
            period_start=None,
            period_end=None,
            trigger_type="scheduled",
            check_only=False,
            force_regenerate=False,
            project_root=PROJECT_ROOT,
            lock_dir=tmp_path / "locks",
        )


def test_previous_period_report_loading(tmp_path):
    """上一周期正式报告进入本期研究包（历史连续性），不存在时用状态基线。"""
    from datetime import date as _date

    from app.assessment.research_driven.generation import _load_previous_period_report

    store = ReportRunStore(tmp_path)
    prev_key = "tainan_mayoral_2026__20260701__20260715"
    store.save(
        {
            "run_key": prev_key,
            "run_id": "rd_prev",
            "period_start": "2026-07-01",
            "period_end": "2026-07-15",
            "generation_status": "ready_for_review",
        }
    )
    prev_dir = tmp_path / "periods" / "20260701_20260715"
    prev_dir.mkdir(parents=True)
    plan = {
        "primary_thesis": {"judgment": "上一期主判断"},
        "trend_outlook": {
            "short_term": "上一期短期判断",
            "key_turning_conditions": ["指标X"],
        },
        "camp_status": [{"camp": "陈亭妃阵营", "status_change": "strengthened"}],
    }
    (prev_dir / "analysis_plan.json").write_text(
        json.dumps(plan, ensure_ascii=False), encoding="utf-8"
    )
    (prev_dir / "final_article.md").write_text("上一期文章正文", encoding="utf-8")

    prev_report, prev_article = _load_previous_period_report(
        store, "tainan_mayoral_2026", _date(2026, 7, 16), _date(2026, 7, 31)
    )
    assert prev_report is not None
    assert prev_report["analysis_plan"]["primary_thesis"]["judgment"] == "上一期主判断"
    assert prev_article == "上一期文章正文"

    # 无上一周期报告 → (None, None)
    empty_store = ReportRunStore(tmp_path / "empty")
    assert _load_previous_period_report(
        empty_store, "tainan_mayoral_2026", _date(2026, 7, 16), _date(2026, 7, 31)
    ) == (None, None)
