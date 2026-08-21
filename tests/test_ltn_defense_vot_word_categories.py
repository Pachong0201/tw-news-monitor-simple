from datetime import datetime, timedelta
from pathlib import Path

import httpx
import pytest
import yaml
from docx import Document

from app.category_classifier import apply_content_classification, classify_content_category
from app.collectors.ltn import LtnRSSCollector
from app.collectors.rss import RSSCollector
from app.database import Database
from app.digest import build_digest
from app.freshness import filter_fresh_articles
from app.importance import load_rules, score_article
from app.main import (
    COLLECTOR_MAP,
    collect_all,
    deduplicate_articles_by_identity,
    deduplicate_articles_by_url,
    validate_sources_config,
)
from app.models import Article
from app.time_utils import TAIPEI
from app.word_digest import build_word_digest


ROOT = Path(__file__).resolve().parent.parent
FIXTURES = ROOT / "tests" / "fixtures"


class StaticClient:
    def __init__(self, response=None, error=None):
        self.response = response
        self.error = error

    def get(self, _url):
        if self.error:
            raise self.error
        return self.response

    def close(self):
        return None


def response(url: str, body: bytes, status: int = 200, content_type: str = "application/rss+xml"):
    return httpx.Response(
        status,
        content=body,
        headers={"content-type": content_type},
        request=httpx.Request("GET", url),
    )


def source(
    source_id: str,
    name: str,
    category: str | None = None,
    url: str = "",
    source_type: str = "rss",
    default_category: str | None = None,
    enabled: bool = True,
):
    cfg = {
        "id": source_id,
        "name": name,
        "type": source_type,
        "url": url,
        "enabled": enabled,
    }
    if category:
        cfg["category"] = category
    if default_category:
        cfg["default_category"] = default_category
    return cfg


def collect_fixture(collector, fixture_name: str):
    body = (FIXTURES / fixture_name).read_bytes()
    collector._client = StaticClient(response(collector.url, body))
    return collector.collect()


def article(title: str, category: str, url: str, published_at: datetime, source_id: str = "media"):
    return Article(
        source_id=source_id,
        source_name="测试媒体",
        category=category,
        title=title,
        url=url,
        published_at=published_at,
        fetched_at=published_at,
        position=1,
    )


def test_production_sources_enable_ltn_defense_and_content_classified_vot():
    config = yaml.safe_load((ROOT / "config" / "sources.yaml").read_text(encoding="utf-8"))
    sources = {item["id"]: item for item in config["sources"]}

    ltn = sources["ltn_defense"]
    assert ltn == {
        "default_category": "military",
        "enabled": True,
        "id": "ltn_defense",
        "name": "自由时报·军武",
        "type": "ltn_rss",
        "url": "https://news.ltn.com.tw/rss/def.xml",
    }

    vot = sources["voice_of_tibet_cn"]
    assert vot["enabled"] is True
    assert vot["url"] == "https://cn.vot.org/feed/"
    assert "category" not in vot
    assert "default_category" not in vot


def test_config_validation_accepts_only_the_two_new_categories_in_addition_to_existing():
    sources = [
        source("mil", "军武", "military", "https://example.com/mil.xml"),
        source("rel", "宗教", "religion", "https://example.com/rel.xml"),
    ]
    validate_sources_config(sources, COLLECTOR_MAP)

    bad = source("third", "第三类", "society", "https://example.com/third.xml")
    with pytest.raises(SystemExit):
        validate_sources_config([bad], COLLECTOR_MAP)


def test_config_validation_accepts_default_category_and_content_classified_sources():
    sources = [
        source(
            "ltn",
            "自由时报·军武",
            url="https://example.com/def.xml",
            source_type="ltn_rss",
            default_category="military",
        ),
        source(
            "vot",
            "西藏之声",
            url="https://example.com/feed/",
        ),
    ]
    validate_sources_config(sources, COLLECTOR_MAP)


def test_ltn_defense_fixture_parses_and_normalizes_complete_article_fields():
    cfg = source(
        "ltn_defense",
        "自由时报·军武",
        url="https://news.ltn.com.tw/rss/def.xml",
        source_type="ltn_rss",
        default_category="military",
    )
    items = collect_fixture(LtnRSSCollector(cfg), "ltn_defense_feed.xml")

    assert len(items) == 3
    first = items[0]
    assert first.source_id == "ltn_defense"
    assert first.source_name == "自由时报·军武"
    assert first.category == "military"
    assert first.title == "国军举行年度联合演训"
    assert first.url == "https://news.ltn.com.tw/news/def/breakingnews/5000001"
    assert first.published_at == datetime(2026, 8, 10, 19, 44, 47, tzinfo=TAIPEI)
    assert first.summary == "演训于台湾东部举行，验证联合指挥能力。"
    assert "<p>" not in first.summary
    assert "<img" not in first.summary


def test_ltn_defense_cross_channel_permanent_url_is_deduplicated():
    cfg = source(
        "ltn_defense",
        "自由时报·军武",
        url="https://news.ltn.com.tw/rss/def.xml",
        source_type="ltn_rss",
        default_category="military",
    )
    defense_items = collect_fixture(LtnRSSCollector(cfg), "ltn_defense_feed.xml")
    duplicate = article(
        "政治频道同文",
        "politics",
        "https://news.ltn.com.tw/news/politics/breakingnews/1234567",
        datetime(2026, 8, 10, 17, 0, tzinfo=TAIPEI),
        source_id="ltn_politics",
    )

    url_unique, url_dups = deduplicate_articles_by_url([duplicate, *defense_items])
    identity_unique, identity_dups = deduplicate_articles_by_identity(url_unique)
    assert len(url_dups) == 1
    assert identity_dups == []
    assert len(identity_unique) == 3


def test_vot_fixture_preserves_unicode_cleans_html_and_classifies_by_content():
    cfg = source(
        "voice_of_tibet_cn",
        "西藏之声",
        url="https://cn.vot.org/feed/",
    )
    items = collect_fixture(RSSCollector(cfg), "vot_cn_feed.xml")
    items = apply_content_classification(items, cfg)

    assert len(items) == 5
    assert {item.category for item in items} == {
        "religion",
        "international",
        "military",
        "politics",
    }
    assert {item.title: item.category for item in items} == {
        "某寺院举行藏传佛教法会": "religion",
        "达赖喇嘛会见美国议员讨论西藏政策": "international",
        "美国国会通过涉藏法案": "international",
        "中印边境解放军举行军事演练": "military",
        "西藏地方政治会议讨论治理政策": "politics",
    }
    assert [item.title for item in items] == [
        "某寺院举行藏传佛教法会",
        "达赖喇嘛会见美国议员讨论西藏政策",
        "美国国会通过涉藏法案",
        "中印边境解放军举行军事演练",
        "西藏地方政治会议讨论治理政策",
    ]
    assert "བོད་" in items[0].summary
    assert "<p>" not in items[0].summary
    assert "&amp;" not in items[1].summary
    assert "John Smith" in items[1].summary
    assert all(item.published_at and item.published_at.tzinfo == TAIPEI for item in items)


def test_vot_content_classifier_cases():
    cases = [
        ("某寺院举行藏传佛教法会", "僧侣与信众参加法会", "religion"),
        ("西藏僧侣宗教活动受到限制", "宗教活动管制", "religion"),
        ("达赖喇嘛主持佛教法会", "", "religion"),
        ("达赖喇嘛会见美国议员讨论西藏政策", "美国议员参加会谈", "international"),
        ("美国国会通过涉藏法案", "法案消息由国会公布", "international"),
        ("中印边境解放军举行军事演练", "报道涉及军事部署与装备", "military"),
        ("西藏地方政治会议讨论治理政策", "会议讨论地方治理", "politics"),
    ]
    for title, summary, expected in cases:
        assert classify_content_category(title, summary) == expected


def test_vot_source_id_is_decoupled_from_religion_category():
    cfg = source(
        "voice_of_tibet_cn",
        "西藏之声",
        url="https://cn.vot.org/feed/",
    )
    items = collect_fixture(RSSCollector(cfg), "vot_cn_feed.xml")
    items = apply_content_classification(items, cfg)
    assert len({a.source_id for a in items}) == 1
    assert len({a.category for a in items}) == 4


def test_vot_stale_fixture_articles_are_not_fresh():
    cfg = source(
        "voice_of_tibet_cn",
        "西藏之声",
        url="https://cn.vot.org/feed/",
    )
    items = collect_fixture(RSSCollector(cfg), "vot_cn_feed.xml")
    items = apply_content_classification(items, cfg)
    now = datetime(2026, 8, 10, 20, 0, tzinfo=TAIPEI)
    result = filter_fresh_articles(items, now)
    assert result.fresh_articles == []
    assert len(result.stale_articles) == 5


def test_content_classification_runs_inside_collect_all(tmp_path, monkeypatch):
    class VotCollector:
        def __init__(self, source_cfg):
            self.source = source_cfg

        def collect(self):
            now = datetime(2026, 8, 10, 12, 0, tzinfo=TAIPEI)
            return [
                article(
                    "某寺院举行藏传佛教法会",
                    "politics",
                    "https://example.com/vot-religion",
                    now,
                    source_id="voice_of_tibet_cn",
                ),
                article(
                    "达赖喇嘛会见美国议员讨论西藏政策",
                    "politics",
                    "https://example.com/vot-international",
                    now,
                    source_id="voice_of_tibet_cn",
                ),
            ]

        def close(self):
            pass

    monkeypatch.setattr("app.main.COLLECTOR_MAP", {"rss": VotCollector})
    db = Database(tmp_path / "classification.db")
    db.connect()
    db.create_tables()
    try:
        result = collect_all(
            [
                source(
                    "voice_of_tibet_cn",
                    "西藏之声",
                    url="https://cn.vot.org/feed/",
                )
            ],
            db,
        )
        inserted = result[0]
        by_url = {a.url: a.category for a in inserted}
        assert by_url["https://example.com/vot-religion"] == "religion"
        assert by_url["https://example.com/vot-international"] == "international"
    finally:
        db.close()


def test_ltn_and_vot_fixtures_are_idempotent_across_two_db_saves(tmp_path):
    ltn_cfg = source(
        "ltn_defense",
        "自由时报·军武",
        url="https://news.ltn.com.tw/rss/def.xml",
        source_type="ltn_rss",
        default_category="military",
    )
    vot_cfg = source(
        "voice_of_tibet_cn",
        "西藏之声",
        url="https://cn.vot.org/feed/",
    )
    ltn_items = collect_fixture(LtnRSSCollector(ltn_cfg), "ltn_defense_feed.xml")
    vot_items = collect_fixture(RSSCollector(vot_cfg), "vot_cn_feed.xml")
    vot_items = apply_content_classification(vot_items, vot_cfg)

    db = Database(tmp_path / "idempotent.db")
    db.connect()
    db.create_tables()
    try:
        first = db.save_articles([*ltn_items, *vot_items])
        second = db.save_articles([*ltn_items, *vot_items])
        assert len(first) == 8
        assert len(second) == 0
        assert db.count_articles() == 8
    finally:
        db.close()


def test_rss_http_429_is_reported_as_source_failure_not_empty_success():
    cfg = source("voice_of_tibet_cn", "西藏之声", url="https://cn.vot.org/feed/")
    collector = RSSCollector(cfg)
    collector._client = StaticClient(
        response(cfg["url"], b"<html>rate limited</html>", status=429, content_type="text/html")
    )
    with pytest.raises(httpx.HTTPStatusError):
        collector.collect()


def test_rss_malformed_xml_without_entries_is_reported_as_parse_failure():
    cfg = source("voice_of_tibet_cn", "西藏之声", url="https://cn.vot.org/feed/")
    collector = RSSCollector(cfg)
    collector._client = StaticClient(response(cfg["url"], b"not an rss document"))
    with pytest.raises(ValueError, match="RSS|feed|parse"):
        collector.collect()


def test_timeout_failure_isolated_and_other_source_still_inserts(tmp_path, monkeypatch):
    class TimeoutCollector:
        def __init__(self, _source):
            pass

        def collect(self):
            raise httpx.TimeoutException("timed out")

        def close(self):
            pass

    class GoodCollector:
        def __init__(self, source_cfg):
            self.source = source_cfg

        def collect(self):
            now = datetime(2026, 8, 10, 12, 0, tzinfo=TAIPEI)
            return [article("正常来源文章", "military", "https://example.com/good", now)]

        def close(self):
            pass

    monkeypatch.setattr("app.main.COLLECTOR_MAP", {"bad": TimeoutCollector, "good": GoodCollector})
    db = Database(tmp_path / "isolation.db")
    db.connect()
    db.create_tables()
    try:
        result = collect_all(
            [
                source("bad", "失败源", "religion", "https://example.com/bad", "bad"),
                source("good", "正常源", "military", "https://example.com/good", "good"),
            ],
            db,
        )
        inserted, total, _dup, failed = result[:4]
        assert total == 1
        assert len(inserted) == 1
        assert failed == ["bad"]
        assert db.count_articles() == 1
    finally:
        db.close()


def test_429_failure_isolated_and_other_source_still_inserts(tmp_path, monkeypatch):
    class RateLimitedCollector:
        def __init__(self, _source):
            pass

        def collect(self):
            raise httpx.HTTPStatusError(
                "429 Too Many Requests",
                request=httpx.Request("GET", "https://example.com/bad"),
                response=httpx.Response(429, request=httpx.Request("GET", "https://example.com/bad")),
            )

        def close(self):
            pass

    class GoodCollector:
        def __init__(self, source_cfg):
            self.source = source_cfg

        def collect(self):
            now = datetime(2026, 8, 10, 12, 0, tzinfo=TAIPEI)
            return [article("正常来源文章", "military", "https://example.com/good", now)]

        def close(self):
            pass

    monkeypatch.setattr("app.main.COLLECTOR_MAP", {"bad": RateLimitedCollector, "good": GoodCollector})
    db = Database(tmp_path / "isolation429.db")
    db.connect()
    db.create_tables()
    try:
        result = collect_all(
            [
                source("bad", "失败源", "religion", "https://example.com/bad", "bad"),
                source("good", "正常源", "military", "https://example.com/good", "good"),
            ],
            db,
        )
        inserted, total, _dup, failed = result[:4]
        assert total == 1
        assert len(inserted) == 1
        assert failed == ["bad"]
        assert db.count_articles() == 1
    finally:
        db.close()


def test_new_categories_do_not_create_automatic_importance():
    rules = load_rules(ROOT / "config" / "importance_rules.yaml")
    for category in ("military", "religion"):
        result = score_article(
            "一般文化交流消息",
            "测试媒体",
            category,
            "没有命中任何重要性规则的普通说明。",
            rules,
        )
        assert result.score == 0
        assert result.level == "normal"
        assert result.matched_rules == []


def test_text_digest_includes_military_and_religion_categories():
    now = datetime(2026, 8, 10, 20, 0)
    digest = build_digest(
        [
            article("军武摘要新闻", "military", "https://example.com/m", now),
            article("宗教摘要新闻", "religion", "https://example.com/r", now),
        ],
        now,
    )
    assert "军武" in digest
    assert "宗教" in digest
    assert "军武摘要新闻" in digest
    assert "宗教摘要新闻" in digest


def test_word_renders_new_media_sections_in_production_order_with_same_style(tmp_path):
    now = datetime(2026, 8, 10, 20, 0)
    items = [
        article("政治", "politics", "https://example.com/p", now),
        article("经济", "economy", "https://example.com/e", now),
        article("军武", "military", "https://example.com/m", now),
        article("国际", "international", "https://example.com/i", now),
        article("宗教", "religion", "https://example.com/r", now),
    ]
    output = build_word_digest(items, tmp_path, generated_at=now)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    headings = ["（一）政治新闻", "（二）经济新闻", "（三）军武", "（四）国际新闻", "（五）宗教"]
    indices = [texts.index(text) for text in headings]
    assert indices == sorted(indices)
    styles = [doc.paragraphs[texts.index(text)].style.style_id for text in headings]
    assert len(set(styles)) == 1


def test_word_section_numbering_remains_contiguous_when_categories_are_empty(tmp_path):
    now = datetime(2026, 8, 10, 20, 0)
    items = [
        article("政治", "politics", "https://example.com/p", now),
        article("国际", "international", "https://example.com/i", now),
        article("宗教", "religion", "https://example.com/r", now),
    ]
    output = build_word_digest(items, tmp_path, generated_at=now)
    texts = [p.text for p in Document(output).paragraphs]
    assert "（一）政治新闻" in texts
    assert "（二）国际新闻" in texts
    assert "（三）宗教" in texts
    assert not any("军武" in text for text in texts)


def test_word_military_articles_sort_newest_first(tmp_path):
    older = datetime(2026, 8, 10, 10, 0)
    newer = older + timedelta(hours=2)
    items = [
        article("较早军武", "military", "https://example.com/old", older),
        article("较新军武", "military", "https://example.com/new", newer),
    ]
    output = build_word_digest(items, tmp_path, generated_at=newer)
    texts = [p.text for p in Document(output).paragraphs]
    assert next(i for i, text in enumerate(texts) if "较新军武" in text) < next(
        i for i, text in enumerate(texts) if "较早军武" in text
    )


def test_official_military_article_stays_in_official_source_section(tmp_path):
    now = datetime(2026, 8, 10, 20, 0)
    official = article(
        "国防部军事新闻",
        "military",
        "https://example.com/official",
        now,
        source_id="mnd_press",
    )
    media = article("媒体军武新闻", "military", "https://example.com/media", now)
    output = build_word_digest([official, media], tmp_path, generated_at=now)
    texts = [p.text for p in Document(output).paragraphs]

    assert texts.index("一、官方信源") < texts.index("二、新闻媒体")
    assert texts.index("一）台湾国防部") < texts.index("二、新闻媒体")
    assert sum("国防部军事新闻" in text for text in texts) == 1
    assert sum("媒体军武新闻" in text for text in texts) == 1
